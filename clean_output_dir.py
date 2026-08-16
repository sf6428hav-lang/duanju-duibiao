import os, sys, re

sys.stdout.reconfigure(encoding='utf-8')
output_dir = r'C:\Users\Administrator\Desktop\短剧对标\output'

cleaned_count = 0
deleted_count = 0

if os.path.exists(output_dir):
    for root, dirs, files in os.walk(output_dir):
        for f in files:
            fp = os.path.join(root, f)
            if f.endswith('.md'):
                with open(fp, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                
                # Strip thinking, TEMPLATEJSON, chat options
                text = re.sub(r'<thinking>[\s\S]*?(?:</thinking>|$)', '', content, flags=re.DOTALL | re.IGNORECASE).strip()
                text = re.sub(r'<think>[\s\S]*?(?:</think>|$)', '', text, flags=re.DOTALL | re.IGNORECASE).strip()
                text = re.sub(r'\[TEMPLATEJSON\][\s\S]*?(?:\[/TEMPLATEJSON\]|$)', '', text, flags=re.DOTALL | re.IGNORECASE).strip()
                text = re.sub(r'```json[\s\S]*?```', '', text, flags=re.DOTALL | re.IGNORECASE).strip()

                has_formal_header = bool(re.search(r'^(?:#+|【)?(?:第[一二三四五六七八九十0-9]+集|第一集|第1集|人物小传|角色设定|人设小传|故事大纲|前十集集纲|对标拆解分析方案)(?:】|\s|$)', text, re.MULTILINE))
                
                lines = text.splitlines()
                clean_lines = []
                in_body = False
                for line in lines:
                    stripped = line.strip()
                    if not in_body:
                        if re.search(r'^(?:#+|【)?(?:第[0-9一二三四五六七八九十]+集|第一集|第1集|人物小传|角色小传|故事大纲|前十集集纲|对标拆解分析方案|一、|二、|三、|▲|场次：)', stripped):
                            in_body = True
                    if in_body:
                        if any(stripped.startswith(k) for k in ["数一下字数", "符合要求", "请确认是否", "你定一下", "请在下方说明"]):
                            continue
                        clean_lines.append(line)
                
                cleaned_text = "\n".join(clean_lines).strip()

                if not has_formal_header or len(cleaned_text) < 100:
                    # Junk file, delete .md and .docx
                    try:
                        os.remove(fp)
                        docx_p = fp.replace('.md', '.docx')
                        if os.path.exists(docx_p):
                            os.remove(docx_p)
                        deleted_count += 1
                        print(f"DELETED JUNK: {f}")
                    except Exception as e:
                        pass
                else:
                    # Clean the file on disk so it contains 100% clean body
                    try:
                        with open(fp, 'w', encoding='utf-8') as file:
                            file.write(cleaned_text)
                        
                        import docx
                        doc = docx.Document()
                        safe_title = f.rsplit('.', 1)[0]
                        doc.add_heading(safe_title, 0)
                        for p in cleaned_text.split('\n'):
                            if p.strip():
                                doc.add_paragraph(p.strip())
                        docx_p = fp.replace('.md', '.docx')
                        doc.save(docx_p)
                        cleaned_count += 1
                        print(f"CLEANED ASSET: {f}")
                    except Exception as e:
                        pass

print(f"Cleanup done. Deleted: {deleted_count}, Cleaned: {cleaned_count}")
