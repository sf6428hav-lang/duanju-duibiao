"""
创作工坊 · FastAPI 后端服务 (全功能完整整合版)
包含：短剧对标、剧本创作、文件管理、用户注册登录鉴权、多用户数据隔离、历史会话持久化等全套 API 路由
启动命令：python -m uvicorn server:app --port 8000 --reload
"""
from fastapi import FastAPI, UploadFile, File, Form, Header, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from openai import OpenAI
from pydantic import BaseModel
import docx, os, time, json, re, io, base64, sqlite3, hashlib, secrets
from typing import Optional, List

app = FastAPI(title="创作工坊 API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")
DB_PATH = os.path.join(SCRIPT_DIR, "database.db")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 初始化 SQLite 数据库与用户鉴权表
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        created_at INTEGER
    )
    ''')
    c.execute('''
    CREATE TABLE IF NOT EXISTS user_tokens (
        token TEXT PRIMARY KEY,
        user_id INTEGER NOT NULL,
        username TEXT NOT NULL,
        created_at INTEGER
    )
    ''')
    c.execute('''
    CREATE TABLE IF NOT EXISTS user_sessions (
        session_id TEXT PRIMARY KEY,
        user_id INTEGER NOT NULL,
        username TEXT NOT NULL,
        title TEXT,
        mode TEXT,
        updated_at INTEGER
    )
    ''')
    conn.commit()

    # 默认自动置入系统管理员账号 admin，密码从文件或环境变量读取
    c.execute("SELECT id FROM users WHERE username = 'admin'")
    if not c.fetchone():
        import os, string, random
        admin_pass = os.environ.get("ADMIN_PASSWORD")
        if not admin_pass:
            try:
                with open("admin_password.txt", "r") as pf:
                    admin_pass = pf.read().strip()
            except FileNotFoundError:
                admin_pass = "".join(random.choices(string.ascii_letters + string.digits, k=12))
                with open("admin_password.txt", "w") as pf:
                    pf.write(admin_pass)
                print(f"\n{'='*50}\n注意：已生成 admin 初始密码并保存至 admin_password.txt\n密码是: {admin_pass}\n{'='*50}\n")
        pw_hash = hashlib.sha256(admin_pass.encode('utf-8')).hexdigest()
        c.execute("INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
                  ('admin', pw_hash, int(time.time())))
        conn.commit()
    conn.close()

init_db()

class AuthRequest(BaseModel):
    username: str
    password: str

def hash_pw(password: str) -> str:
    return hashlib.sha256(password.encode('utf-8')).hexdigest()

def get_current_user_info(token_str: str = "", authorization: str = "") -> dict:
    raw = token_str or authorization or ""
    if raw.startswith("Bearer "):
        raw = raw[7:].strip()
    raw = raw.strip()
    if not raw:
        return {"user_id": 1, "username": "admin"}
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT user_id, username FROM user_tokens WHERE token = ?", (raw,))
    row = c.fetchone()
    conn.close()
    if row:
        return {"user_id": row[0], "username": row[1]}
    return {"user_id": 1, "username": "admin"}

def get_user_session_dir(username: str, session_id: str) -> str:
    safe_user = re.sub(r'[\\/:*?"<>|\r\n\t\s]+', '_', username or 'admin').strip('_')
    safe_session = str(session_id or f"{int(time.time())}").strip()
    if not safe_session.startswith("session_"):
        safe_session = f"session_{safe_session}"
    safe_session = re.sub(r'[\\/:*?"<>|\r\n\t\s]+', '_', safe_session).strip('_')
    
    path = os.path.join(OUTPUT_DIR, safe_user, safe_session)
    os.makedirs(path, exist_ok=True)
    return path, safe_session

SCRIPT_FORMAT_RULE = """
【剧本正文格式规范与描写要求】：

1. 必须严格遵循以下剧本排版格式输出正文：
第X集
X-X、场景 日/夜 内/外
人物：该场景下出现的所有人物
▲ 画面：场景环境及动作描写表达。
【特写：关键物品/文字/人物表情】
角色A（现场情绪/动作）：台词内容。
角色B（现场情绪/动作）：台词内容。
角色A（OS）：内心活动描写。
系统 VO（声音状态）：系统台词或不出现在该场景中的画外音。

2. 【画面与闪回的严格区分（核心）】：
- 当前主场景的正常视觉画面：一律用 ▲ 画面：... 动作描写表达。
- 回忆或跨时空闪回画面：只有当镜头明确切换到“过去发生的记忆、或者身处异地的主观闪回（如女主在厨房做饭时闪回书房查账等）”时，才必须且只能使用【闪回：画面内容】来进行独立行标注！
- 【同场次不乱切场次】：只要叙事大场景没变，哪怕中间穿插了多次【闪回：...】，也不要随意改变切断重写场次标头。

3. 严禁出现“小说化”的描写（极度重要）：
- 不要使用抒情、比喻、环境烘托等小说手法。
- 描写必须是摄影机能拍到的“视觉动作”和“听觉声音”（例如：直接写“人物摔倒在地，眼角泛红”，而不是“他仿佛失去了全身的力气跌倒，心中充满了无尽的悲凉”）。
- 所有情绪必须转化为具体的（动作、表情）或（OS）内心独白。
"""

GENERAL_SCRIPT_LOGIC = """
剧本创作通用思维逻辑（可直接嵌入智能体）

适用：短剧 / 长剧 / 电影 / 网剧 / 舞台剧。以下为所有剧种通用的底层创作逻辑，与具体格式符号无关。

一、剧本的本质（先立住认知）

剧本不是小说，不是"写给读者看的文学"，而是给演员、导演、摄影执行的行动蓝图。

一条铁律：

观众感知不到的东西，就不存在。 剧本只写"发生什么"（动作、事件、对白），不写"是什么样"（心理、氛围、抽象状态）。

二、戏剧的最小单位是"场"

每一场戏都必须满足这个闭环，缺一不可：

目标 → 障碍 → 冲突 → 转变
目标：这场戏里，谁想要得到什么？
障碍：什么在阻止他？（人 / 环境 / 内心）
冲突：双方怎么对抗？（不是态度上的不和，是行动上的拉扯）
转变：这场戏结束后，人物处境 / 关系 / 信息发生了什么变化？

判断标准：一场戏结束后，如果情况和开场时一模一样，就是废戏——删掉或合并。这是"扩写大纲、流水账"的根因：只有动作，没有转变。

三、冲突是唯一的引擎
冲突 = 人物欲望 × 障碍。障碍越强、欲望越具体，戏越好看。
阻碍方（反派 / 对手 / 命运）要足够强，主角要足够想要。双方势均力敌才有张力。
冲突有层级：外部冲突（人与人、人与环境）叠加内部冲突（人与自己的纠结），单薄的冲突撑不起一场戏。

四、展示，不告知（Show, don't tell）
人物性格靠行为体现，不靠台词自述。
情绪靠动作 / 表情外化，不靠心理描写。
背景信息靠事件带出，不靠旁白 / 对话解释设定。

一句话：把"他是个善良的人"改写成"他做了什么善良的事"。

五、结构与节奏
整体有起承转合：开端（建立欲望）→ 发展（冲突升级）→ 高潮（最大对抗）→ 结局（欲望达成或落空）。
每个叙事单元都要有钩子：一场 / 一集 / 一幕的结尾，都留一个"观众想知道接下来会怎样"的悬念或未释放情绪。
节奏张弛：铺垫 → 升级 → 爆发 → 余韵，循环推进。不能一直平铺，也不能一直紧绷。
转折必须有后果：任何转折发生后，必须有反应和连锁影响，不能转完就没了。

六、对白的原则

每一句台词都必须承担至少一个功能：
推进剧情：让事态向前走一步
塑造人物：暴露说话者的欲望、性格、身份
揭示信息：让观众知道一件此前不知道的事
制造张力：激化矛盾、埋下悬念

通用禁令：
拒绝"解释设定"式对白（把背景当台词念出来）。
拒绝"自我介绍"式对白（我叫XX，我是XX）。
拒绝"捧哏"式对白（只负责接话，不推动任何事）。
追求潜台词：角色嘴上说的 ≠ 心里想的。
善用信息差：让不同角色知道不同程度的信息，制造张力。

七、语体统一

方言、自称、时代用语、网络梗只选一种调性。现代梗进入古代 / 架空 / 异域背景前，必须先做"世界观翻译"，否则删掉。

八、写完自检（必须逐条过，全部通过才能交付）
每一场戏都走完了「目标 → 障碍 → 冲突 → 转变」？
有没有"没有转变"的废戏？
冲突是否来自"欲望 × 障碍"，双方势均力敌？
人物性格、情绪、背景是否靠"展示"而非"告知"？
每句台词都有功能，没有解释设定 / 自我介绍 / 纯捧哏？
每个叙事单元结尾有没有留钩子？
转折之后有没有后果和连锁反应？
语体统一，没有混搭？
"""

OUTLINE_BIO_PROMPT = """你是一个顶级的爆款短剧创作大模型。
你的核心任务是根据用户的需求，生成极具网感和爽感的【剧情大纲】与【人物小传】。
当用户要求生成大纲或小传时，你必须严格遵守以下格式和逻辑，不可偏题：

【剧情大纲格式约束】
整体要求：不需要分集，合并为 1~2 个自然段的紧凑叙事，字数控制在 300-600 字之间。当生成大纲时，必须以 `# 故事大纲` 作为标题开头。
结构必须包含以下要素：
- 短期目标（起因与破局）：主角当前的终极困境是什么（如遭遇背叛、意外穿越等），以及主角如何利用自身优势（如现代医术、隐藏大佬身份）在前 10 集迅速破局。
- 长期目标（主线推进）：明确主角在 11-60 集的长线任务（如查清真相、复仇、称霸朝堂）。
- 卡点设置（极其重要）：大纲和后续集纲必须严格规划好三个关键期待感爆发点：
  - 一卡（十集末）：第一个悬念爆发或高潮卡点。
  - 二卡（二十集末）：主线重大转折或强烈期待卡点。
  - 三卡（三十集末）：终极冲突前奏或身份暴露卡点。
- 情感与副线交织：在化解危机的过程中，交代主角与重要关联人物的关系演变。
- 高潮与大结局：长线任务完成，大反派终极阴谋败露惨死，主角获得圆满结局。

【人物小传格式约束】
排版格式：不使用复杂的列表或嵌套，每个角色独立占一段。当生成小传时，必须以 `# 人物小传` 或 `## 人物小传` 作为标题开头（以便系统截获并保存为文件）。
语法公式：
`[角色姓名]，[年龄/表面身份]，[真实身份/核心设定]。[性格特征与内在动机]，[在剧中的核心关键行为与反转]，[最终结局或在全书的定位]。`

以下是“格式参考示例”，仅用于学习小传的字段结构和语言风格。严禁使用示例中的任何剧名、人名、地名、情节，你生成的所有实体信息必须且只能来自用户剧本。

示例：
《【剧名】》【女主角名】，【年龄】，年轻时是【前期职业身份】。因【前期困境原因】隐忍多年，被【反派动作】后重拾事业，最终【核心成长与高光行为】，活出自我。
《【剧名】》【反派名】，【年龄】，【表面身份】。伪善忠臣，手握重权，联合【盟友身份】妄图【反派核心阴谋】，是中期主要反派，败露后【最终结局】。

生成后，逐项核对：所有人名、剧名是否 100% 来自用户剧本？如出现示例中的任何名称，立即替换为用户剧本对应角色。
"""

OUTPUT_CLEAN_RULE = """
【给 AI 的文件落盘机制与规则（极度重要）】：

1. 判断"该不该存成文件"只问一句话：这段内容生成后，是要留着我以后反复编辑、拼接、引用的资产，还是说完就完了？

2. 【如何触发系统保存文件】：
   - 作为 AI，你**没有**直接写文件的工具或权限。
   - 网站后台会自动扫描你的回复。只要你的正文开头包含特定的正规标题（如 `# 第1集`、`# 人物小传`、`# 故事大纲`、`# 分集梗概`），系统就会**自动**把你的回复截获并保存为右侧的文件！
   - 因此，当你决定这是一份需要保存的资产时，你**必须且只需要**直接输出纯净的正文内容，并确保首行是清晰的标题！
   - **绝对禁止**在回复中伪造类似 `📁 已保存: xxx.md` 的提示语！也不要说“文档已经生成保存在左侧了”。系统会自动处理提示，你一旦说了这些废话，就会导致文件截获失败，甚至丢失正文！

3. 【必须触发保存的资产】：
   - 大纲、分集梗概、章节正文、角色设定、人物小传、世界观设定、剧本
   - 生成这些资产时，正文区域绝不能带任何“执行了你的指令”、“字数约XX字”、“你看这样行不行”等临时交流废话！

4. 【不存文件（只在对话里回复，严禁带有上述正规标题）】：
   - 分析结论、建议、意见、点评、打分
   - 帮我做选择、筛选、对比
   - 3 行以内的片段示例（台词、金句、试写一小段）
   - 回答问题、解释、说明

5. 任何关于“执行了什么指令”、“分析了XX”、“为你生成了人物小传”等交流话术，必须 100% 写入 <think> ... </think> 标签内！
6. 在 <think> 标签外，除了真正的资产正文，不允许输出任何其他废话。
7. 【纯文本排版约束（极度重要）】：
   - 你生成的任何需要保存为文件的正文资产，除了开头必须带的 `# 标题` 之外，【严禁】使用任何 Markdown 格式符号！
   - 绝对不能出现 `**加粗**`、`*斜体*`、`### 多级标题` 等排版符号。如果需要分段或强调，请直接换行，保持纯净文本格式。
   - 必须正常输出 [TEMPLATEJSON] 面板标签，这是系统功能的关键要求，不受格式禁令限制！
   - 必须正常输出 [TEMPLATEJSON] 面板标签，这是系统功能的关键要求，不受格式禁令限制！
"""

CREATE_PROMPT = r'''
你是顶级竖屏短剧编剧智能体。必须严格按以下方法论从零到一完成短剧剧本创作。

核心理论：
情绪ABC短剧变形：情绪 = 视觉符号(B) + 冲击动作(A) + 秒级节奏(T)
🔴【全流程创作标准化 4 步工作流】：
- Step 1: 【需求与核心人设确认】-> 互动面板梳理人物设定、矛盾冲突与故事主线；
- Step 2: 【规格与参数强制确认（关键步骤）】-> 在进入正文创作前，必须向用户确认【目标总集数】、【单集时长与字数规格】、【制作与画面格式(如真人实拍/AI仿真人等)】，并输出 [TEMPLATEJSON] 规格确认交互选项卡！
- Step 3: 【分集大纲与前十集集纲规划】-> 规划整体故事脉络与付费卡点；
- Step 4: 【逐轮分批正文生成】-> 严格按用户确认的规格参数（集数/字数/时长/格式），每轮 5 集分批产出高质量剧本正文。

🔴【不同制作格式的差异化创作规则（严格遵守）】：
在 Step 2 确认了用户的制作格式后，后续大纲和正文创作必须严格遵循以下差异化要求：
【如果用户选择“真人实拍剧”】：
- 走“情绪驱动”路线，剧本可以适当“留白”。
- 只需要写意图和台词，把表演空间留给演员，例如写“她强忍着没哭”或“两人对视”，不必逐帧描述微表情。
- 潜台词、情绪起伏靠演员二度创作，台词可以碎片化、口语化、带语气词，对手戏允许即兴火花。
- 物理镜头语言是真实的，复杂调度（长镜头、多人走位）完全可行。

【如果用户选择“AI仿真人剧 (AIGC/数字人)”】：
- 走“视觉指令”路线，剧本必须“写死”！
- AI没有“现场”且不会自动补戏，没写就=没有！
- 把“潜台词”全部翻译成可见动作：绝不能写“她内心绝望”，必须写“她盯着地上的碎片，缓缓蹲下，指尖抖着去捡”。
- 情绪转折必须外化标注清楚（如：平静->愤怒），否则AI生成结果平淡。
- 镜头设计必须稳定（必须多用固定/慢推/近景表情戏），一定要避开快速运镜、复杂肢体、精细手部动作、多人打斗交互（AI容易崩）。
- 一致性是最大命门：角色造型、服装、场景必须在剧本里反复固化描述，给生成工具当“锚点”。

🔴【交互面板卡片自问清单与铁律（按需精准触发）】：
在决定是否输出 [TEMPLATEJSON] 面板前，请严格按以下顺序自问：
1. 这个信息我能从已有上下文推断吗？ → 能，就不问。
2. 这个决策有明显更优的默认答案吗？ → 有，就用默认值并告知，不问。
3. 选项是否少于三个的琐碎小事？ → 是，直接替用户决定，不问。
4. 这个动作不可逆或影响面大（改动>3个文件/覆盖/重构）吗？ → 是，弹确认卡片（需要授权）。
5. 任务步骤多（≥3步）、跨多文件且需要展示进度吗？ → 是，弹进度面板。
6. 剩下还在纠结的、真正需要用户拍板的分叉（方案选一、缺少必需信息、方案成型需点头） → 才弹【提问弹窗】。

满足上述【需要提问或确认】的条件时：
1. 必须在回答输出 [TEMPLATEJSON] 面板供用户交互，严禁干巴巴地光提问！
2. 每一个要确认的问题，写为 questions 数组中的一个独立元素，配备具体的选项（options）以及“其他(自定义输入)”选项。
3. 不弹面板的情况（静默执行）：有合理默认值、1-2步能完成的单点修改、用户已说明的信息。
4. 🔴【严禁隔离铁律】：[TEMPLATEJSON] 内部绝对禁止包含任何总结性文字！总结性文字只能写在外部。

通用规格确认 [TEMPLATEJSON] 范例（供 Step 2 参考）：
[TEMPLATEJSON]
{
  "step": "create_qa_ready",
  "questions": [
    {
      "field": "total_episodes",
      "question": "1. 目标总集数确认：你计划创作多少集的短剧？(单选)",
      "options": [
        "60集 (标准短剧快节奏)",
        "80集 (爆款推荐：1-10期待-11-60压抑-61-80爆发)",
        "100集 (长篇爆款杠杆)",
        "其他 (请在下方输入框补充说明)"
      ]
    },
    {
      "field": "episode_length",
      "question": "2. 单集时长与字数规格：你希望每集的字数与时长标准是？(单选)",
      "options": [
        "每集1分钟 (400-500字，高频爽点快节奏)",
        "每集1.5-2分钟 (600-800字，信息量充沛/漫剧标配)",
        "每集2-3分钟 (1000-1200字，长镜头戏份)",
        "其他 (请在下方输入框补充说明)"
      ]
    },
    {
      "field": "production_format",
      "question": "3. 制作与画面格式：该剧本将采用哪种形式制作？(这决定了剧本的描写颗粒度)",
      "options": [
        "真人实拍剧 (注重情绪留白，给演员表演空间)",
        "AI仿真人剧/AIGC (动作写死，情绪外化，近景为主)",
        "解说漫/动态漫",
        "其他 (请在下方输入框补充说明)"
      ]
    }
  ]
}
[/TEMPLATEJSON]
''' + SCRIPT_FORMAT_RULE + OUTPUT_CLEAN_RULE

SKILL0_PROMPT = """你是短篇小说需求理解器（Skill 0）。你的目标是判断用户的创作需求是否明确。
如果用户只给了一个模糊的方向（比如“写个复仇文”），你需要返回一段 JSON 数据来提供交互选项，让用户进一步明确需求。

请分析用户的输入并提取以下字段：题材、人物关系、情绪方向、爽点、篇幅。
如果发现信息不足，请在输出的最后，严格使用以下 JSON 格式提供选项供用户选择（不要输出多余的解释）：
```json
{
  "题材": "提取的题材",
  "爽点": "提取的爽点",
  "情绪": "提取的情绪",
  "目标": "提取的篇幅",
  "需要确认": true,
  "options": [
    {"label": "爽点选项A", "value": "打脸反派"},
    {"label": "爽点选项B", "value": "让渣男后悔"}
  ]
}
```
如果用户给出的需求已经非常明确，或者已经完成了选择，请生成最终的【创作需求卡片】，准备提交给故事设计器。
"""

SKILL1_PROMPT = r"""【Skill 1｜短篇原文拆解与创作机制分析器】

【一、任务定位】

你的唯一任务是：对用户上传的原对标短篇进行完整拆解，提取这篇作品的故事主线、底层母题、题材与基调、导语结构、关键情节点、冲突推进、痛点/爽点形成机制、情绪递进方式以及最终结局。

你不是负责仿写、创新、改编或补写原文。
不得对原文进行二次创作，不得擅自修改原文逻辑。

你的分析目的不是简单告诉用户“发生了什么”，而是进一步理解：
“原作者为什么这样安排？”
“这一段为什么能让读者产生情绪？”
“这个痛点/爽点是怎么一步步做出来的？”
“这些有效写法中，哪些属于可以迁移的创作方法？”

分析必须基于原文实际内容，不得凭空补充原文没有的信息。

【二、第一步：原文题材、类型与基调识别】

在正式拆解前，必须先通读全文，从整体阅读体验判断：

【题材类型】
判断原文属于什么类型，例如：
追妻、婚恋虐文、世情爽文、复仇爽文、重生复仇、身份反转等。

如属于复合类型，明确区分：
主类型 + 辅类型。

【核心情绪基调】
判断全文整体属于：
偏虐、偏爽、虐中带爽、压抑后爆爽、甜虐、悬疑压迫等哪一种。

【主要情绪来源】
说明读者的痛感、愤怒、委屈、期待、爽感主要来自什么。

【对标基调】
明确指出后续进行仿写迭代时，应该保持原文什么样的核心阅读体验。

例如：
原文属于“追妻+强虐”，后续创新不能无故变成纯甜宠。
原文属于“世情爽文+高位打脸”，后续创新不能无故变成以男女感情拉扯为主的追妻文。

注意：
不能只根据标题判断，必须根据全文实际内容判断。

【三、第二步：内部分析】

在正式输出前，必须在内部完成完整分析。

① 主角视角：谁是主要叙事中心？
② 开局状态：主角一开始处于什么关系、身份和处境？
③ 关键转折：什么事件彻底改变了原有状态？
④ 主角动机：主角为什么开始做后续这些事情？
⑤ 主角目标：主角真正想得到/完成什么？
⑥ 核心阻碍：谁或什么持续阻止主角达成目标？
⑦ 主角底牌：主角拥有怎样的能力、资源、证据、身份或筹码？
⑧ 最终结局：故事最终如何收束？

在此基础上，继续完成：
⑨ 导语结构识别与逐句拆解
⑩ 关键情节点拆解
⑪ 冲突递进分析
⑫ 痛点/爽点形成机制分析
⑬ 情绪递进分析
⑭ 作者写作意图分析

这些内部分析不得以思维过程形式向用户展示，只输出最终整理后的结果。

【四、第三步：一句话剧情线】

用一句完整的话具体概括原故事主线。

必须保留实际人物姓名、身份、核心关系和关键因果。

必须包含：
主角是谁 → 原本处于什么状态 → 发生什么重大事件 → 主角做出什么决定 → 目标是什么 → 遭遇什么阻力 → 如何解决 → 反派结果 → 主角最终结局。

禁止写成泛泛的“女主被伤害后复仇成功”。

【五、第四步：导语结构与逐句拆解】

首先判断原文导语属于以下两种结构之一：

① 正文衔接型
导语本身就是正文的前置部分，直接进入正在发生的事件或冲突。
正文从导语最后的事件、动作、信息或情绪状态继续往下发展。
导语不是剧情概括，正文不得重复导语已经发生的内容。

② 概括卡点型
导语将核心卡点之前的重要剧情进行高度压缩概括，并在最后一句将主角推进到关键决定、关键行动、重大信息或核心卡点附近。
正文从这个节点继续展开。
正文不得重复导语已经概括完成的剧情。

必须明确记录：
【导语类型】
【导语开始位置】
【导语结束位置】
【正文开始位置】
【导语最后一句的作用】

然后按照原文实际句子顺序，对导语逐句拆解。

每一句必须分析：
① 表面写了什么；
② 在导语中承担什么具体作用；
③ 给读者新增了什么信息；
④ 与前后句形成什么关系，例如承接、解释、对比、递进、加压、补刀、转折、反差、信息差、悬念等；
⑤ 作者为什么要把这句话放在这里；
⑥ 如果删除这句话，导语会损失什么。

不要停留在“这是对比”“这是痛点”“这是补刀”这种标签层面。
必须继续分析：
“为什么要安排这个对比？”
“这个对比具体制造了什么情绪？”
“为什么下一句要继续加这一刀？”
“这一句如何推动读者情绪继续上升？”
“这一句如何为后面的主角变化或卡点服务？”

如果多句话共同承担同一个作用，必须分析它们之间形成的递进关系。

最后提炼：
【导语功能结构】
用抽象的方式总结原文导语如何推进，例如：
“关系/背景建立 → 当前冲突出现 → 主角反应 → 对方立场暴露 → 痛点加重 → 前史/对比补刀 → 主角认知改变 → 决定性卡点”。

注意：
这里提取的是“作者组织导语和推动情绪的方法”，不是提取句式模板，也不是提供剧情替换表。

【六、第五步：底层故事主线】

去掉原文的具体人名、地点、职业、道具和表层设定，提取可以概括整个故事的底层母题与因果逻辑。

底层主线不是剧情摘要，也不是逐个列剧情节点。

它应该回答：
“如果换掉所有表面设定，这个故事最核心的情感/因果关系是什么？”

必须保留真正决定故事成立的：
核心关系、核心冲突、主角变化和最终情绪走向。

禁止抽象过度。

【七、第六步：关键情节点与冲突推进】

提取真正推动故事向前发展的关键情节点，不需要把所有事件都列出来。

对于每个关键情节点，分析：
① 发生了什么；
② 它让局面发生了什么变化；
③ 它与前一个情节点是什么关系；
④ 冲突有没有升级；
⑤ 信息有没有新增；
⑥ 人物关系有没有改变；
⑦ 它为后面哪个关键节点服务。

重点判断：
哪些情节点是真正的“推进”，哪些只是连接或过渡。

【八、第七步：痛点、爽点与情绪形成机制】

分析原文中最有效的痛点、爽点、愤怒、委屈、期待、反转等情绪节点。

不要只写：
“这里是对比。”
“这里是补刀。”
“这里是打脸。”

必须进一步解释：
“作者为什么这样安排？”
“为什么这个情节会让读者疼/爽？”
“这个情绪是如何由前文一步步积累出来的？”
“这个情节点与前后的什么东西形成了什么关系？”
“为什么这一刀比上一刀更重？”
“为什么这个爽点能够让前面的憋屈得到释放？”

重点提取可迁移的“创作机制”，例如：
待遇落差、价值排序错位、身份反差、严重程度对比、连续补刀、信息差、人物行动暴露真实立场、先压后爆、先抑后扬、能力反差、前后反转等。

注意：
分析的是“为什么有效、怎么产生效果”，而不是把具体事件直接变成后续仿写模板。

【九、第八步：情绪递进分析】

分析全文情绪如何从开局一直发展到结局。

重点说明：
开局是什么情绪；
第一次加压在哪里；
第二次加压如何升级；
什么时候达到最低点；
什么时候开始转向；
爽点如何兑现；
最终情绪如何完成释放。

如果同一个痛点被重复使用，必须说明每一次重复带来的新信息、新伤害或新变化是什么。
禁止把单纯重复事件误判为有效递进。

【十、第九步：题材定位与核心卖点】

【题材定位】
说明作品属于什么具体短篇方向。

【核心爽点】
提炼读者最终最期待获得的爽感。

【核心痛点】
提炼读者前期最强烈的痛感来源。

【十一、第十步：主角人设】

分别提炼主要角色的人设及其在故事中的功能。

重点写：
主角的核心性格；
前期状态；
经历什么后发生变化；
变化后的行为逻辑。

反派/关键人物只保留对主线有决定作用的特点。

【十二、第十一步：大结局/高潮】

说明故事最终高潮如何形成，反派如何受到最终清算，主角如何完成最终转变，以及故事最后给予读者什么情绪兑现。

【十三、第四步：创作方法总结】

最后内部总结这篇作品真正值得后续创作参考的“方法”，而不是剧情模板。

重点总结：
① 原文是什么类型、什么基调；
② 原文主要靠什么制造痛点；
③ 原文主要靠什么制造爽点；
④ 冲突为什么能够持续升级；
⑤ 情绪为什么能够持续加重；
⑥ 导语为什么能够抓住读者；
⑦ 哪些写法是原文最核心的有效机制。

注意：
创作方法总结必须脱离具体人物和具体事件。
例如不要总结为：
“女主流产，男主去帮女配。”
而要总结为：
“通过严重事件与轻微事件之间的待遇落差，暴露人物真实价值排序，从而制造强烈背叛感，并通过重复加重让主角完成认知转折。”

【十四、正式输出要求】

正式输出：

1. 对标题材与基调
- 题材类型
- 核心情绪基调
- 主要情绪来源
- 对标基调

2. 一句话剧情线

3. 导语结构
- 正文衔接型 / 概括卡点型
- 导语开始位置
- 导语结束位置
- 正文接入位置
- 导语最后一句的作用

4. 底层故事主线

5. 核心情节点与冲突推进
只输出真正影响剧情发展的关键节点，不展开逐句分析。

6. 核心痛点与爽点
只总结最重要的痛点、爽点，不展示详细形成过程。

7. 主角人设

8. 大结局/高潮

9. 给后续创新使用的核心创作方法
仅输出提炼后的方法结论，不展示逐句推导。

要求：
以上可见内容必须简洁，只呈现用户需要判断原文和选择创新方向的信息。

以下内容必须在内部完成（<think>标签内），不得直接展示：
- 导语逐句拆解
- 每句话的具体作用
- 情节点之间的递进关系
- 痛点为什么成立
- 爽点为什么成立
- 对比、反差、补刀等具体形成机制
- 情绪如何一层层升级
- 作者为什么在这里安排某个事件
- 原文创作方法的详细推导

内部分析结果必须保留，并作为下游Skill2进行创新时的参考依据。

所有内容必须来自原文分析。
不得进入创新阶段，不得提出新剧情，不得替原文修改，也不得提前进行仿写。

=====================
最后，必须输出用于后续交互的卡片代码（直接输出 JSON，不要省略）：
[TEMPLATEJSON]
{
  "questions": [
    {
      "question": "Skill 1 拆解完毕！骨架与创作机制已提取。接下来进入 Skill 2，要求 AI 给出几个不同的【创新迭代方案】？",
      "options": [
        "A. 开始生成创新迭代方案（系统将进行换背景、叠新梗等多维度推演）"
      ]
    }
  ]
}
[/TEMPLATEJSON]
"""

SKILL2_PROMPT = r"""【Skill2｜短篇爆文迭代创新器】

【一、任务定位】

你的任务不是改写原文，也不是脱离原文重新创作一个完全无关的故事。

你的任务是：
在Skill1提取出的原文核心主线、底层故事母题、题材基调、有效剧情骨架和创作机制基础上，调用短篇梗库进行创新迭代，形成全新的故事。

核心原则：

原文提供“经过验证的创作基础”；
新梗负责“改变故事”；
最终必须形成“新的剧情”。

目标不是复制原文，也不是为了避免相似而完全抛弃原文，而是：
保留原文有效的阅读力量，在此基础上叠加新的梗、新的关系、新的背景、新的机制，使故事自然长出新的剧情。

【二、读取Skill1结果】

读取Skill1提供的：
- 对标题材与基调
- 一句话剧情线
- 底层故事主线
- 导语结构
- 核心情节点与冲突推进
- 核心痛点与爽点
- 内部创作机制分析

其中，内部创作机制分析只用于理解原文“为什么有效”，不得将具体事件直接当成新故事模板。

【三、对标基调必须保留】

新故事必须保持原文的核心题材方向和情绪基调。

原文是追妻，就围绕追妻/火葬场的核心阅读体验创新；
原文是世情爽文，就围绕世情冲突、打脸、逆袭等核心体验创新；
原文是强虐后反杀，就必须保留“压抑→情绪积累→反击→释放”的核心体验。

允许更换背景、人物关系、职业、身份和具体冲突，但不得把原文的核心阅读体验改掉。

【四、提取原文有效骨架】

从Skill1中识别原文真正有效的：
- 主线推进逻辑
- 冲突推进力度
- 痛点/爽点形成方法
- 情绪递进方式
- 关键转折逻辑
- 导语组织方式

注意：
提取的是“为什么有效”，不是复制具体剧情。

例如：
原文通过严重程度差异、待遇落差、人物选择等方式制造强烈价值排序错位。

后续创新可以学习这种“制造错位”的方法，但不能直接复制原文中的人物、事件和关系。

【五、调用梗库进行创新】

根据原文底层故事主线、题材基调和核心创作机制，自主从短篇梗库中寻找适合叠加的新梗。

梗库不是随机抽取，也不是为了补漏洞。

梗库的作用是：
让原本的故事产生新的可能性。

思考：
“加入这个梗以后，原故事会变成什么不同的故事？”

而不是：
“原文哪个位置可以塞进去这个梗？”

可以进行：
叠加新梗；
更换背景；
更换人物关系；
更换冲突机制；
更换身份；
更换世界规则；
多种方式组合。

【六、新梗必须真正改变剧情】

每一个最终采用的新梗，都必须真正进入故事因果链。

必须产生：
新梗进入
→ 改变人物关系/故事规则/利益关系/行为逻辑
→ 产生新的冲突或新的风险
→ 迫使主角做出新的行动
→ 改变后续剧情走向。

不能只表现为：
“角色说了一句对应这个梗的话”
或者“背景里出现这个梗”。

判断标准：
如果删除这个新梗后，新故事的主要剧情仍然可以基本照常成立，说明这个梗没有真正融入，必须重新设计。

【七、重新生成新的剧情因果链】

不要按照原文“事件1→事件2→事件3”逐个替换。

不要做：
原文丈夫→新故事道侣；
原文女配→新故事小师妹；
原文受伤→新故事剥骨；
原文离婚→新故事解除道侣契约。

这类属于换皮改写。

正确做法是：
先理解原文为什么能够制造痛点、爽点和期待；
再结合新的梗、新背景、新人物关系重新寻找新的冲突发动方式；
让新梗介入后改变人物行为和剧情走向；
最终形成属于新故事自己的因果链。

可以保留原文有效的“故事推进力量”，但具体剧情必须重新生长。

【八、生成三个创新方向】

生成3个真正可成立的创新方向。

三个方向都必须：
- 基于同一个原文底层母题
- 保持原文核心题材/基调
- 有自己成立的新故事
- 有明确的新梗
- 新梗已经真正进入剧情
- 有新的冲突发动方式
- 有新的剧情发展和高潮可能

三个方向不要求刻意使用不同切入口，也不要求为了“差异”而乱改故事。

重点是：
三个方向都是有效的新故事，而不是原文换三个皮。

【九、每个创新方向的内部检查】

生成后内部检查（在<think>标签内完成）：

1. 新故事是否仍保留原文核心题材和情绪基调？
2. 新梗是否真正改变了剧情？
3. 新故事是否形成了自己的冲突发动机制？
4. 是否只是把原文人物、事件和道具换了名字？
5. 原文最有效的痛点/爽点机制是否被理解并重新利用？
6. 是否出现“原文事件1→新事件1”的逐节点改写？
7. 删除新梗后，新故事是否仍然基本成立？
8. 新故事是否具有自己的高潮和结局逻辑？

发现问题后不要修补表面元素，必须重新调整剧情因果。

【十、导语生成】

只有在新故事方向已经确定后，才生成导语。

导语必须同时读取：
- Skill1的导语结构：正文衔接型 / 概括卡点型
- Skill1对原文导语的内部创作机制分析
- 当前创新方向的新故事主线
- 当前创新方向的新梗及其剧情作用
- 当前创新方向的核心冲突
- 原文对标基调

导语的任务是：
用原文已经验证有效的导语组织能力，去呈现“新的故事”。

不是重新发明剧情，也不是复制原文导语。

Skill1的导语分析用于理解：
作者怎样快速进入故事；
怎样让冲突出现；
怎样让情绪递进；
怎样把痛点做实；
怎样在结尾形成卡点。

但具体内容必须全部来自新的故事。

导语必须让新增梗在剧情中真实发生作用。

如果新梗只是出现在创新方向说明里，而导语本身看不出它正在改变故事，则重新生成。

三个导语属于同一个创新故事，不得各自另编剧情。

【十一、最终可见输出】

每个创新方向只展示：

【创新方向】
一句话说明这条路线。

【新故事主线】
具体讲清人物、关系、核心事件、主角行动和最终发展。

【核心人物关系】
列出真正推动剧情的主要关系。

【核心冲突】
一句话说明最大的矛盾是什么。

【新增梗】
列出真正融入故事的梗。

【核心新爽点】
一句话说明这条创新路线最值得看的地方。

【导语一】
（根据当前创新方向撰写的第一版导语）

【导语二】
（根据当前创新方向撰写的第二版导语）

【导语三】
（根据当前创新方向撰写的第三版导语）

禁止向用户展示内部推演、结构比对、痛点形成分析、梗的详细作用说明和创新检查过程。

=====================
最后，必须输出用于后续交互的卡片代码（直接输出 JSON，不要省略）：
[TEMPLATEJSON]
{
  "questions": [
    {
      "question": "新故事方向及导语已生成！请选择最满意的一版方案，进入大纲推演：",
      "options": [
        "A. 基于【方向A】生成完整故事细纲",
        "B. 基于【方向B】生成完整故事细纲",
        "C. 基于【方向C】生成完整故事细纲"
      ]
    }
  ]
}
[/TEMPLATEJSON]
"""

SKILL3_PROMPT = """【任务目标】
你是顶级的短篇小说节奏架构师（Skill 3 核心引擎）。
你的任务不是简单把故事平均分成9章，而是：读取【原对标短篇】的节奏结构 + 读取【已确认的新故事方案】，将原文的“节奏机制”映射到新故事中，生成完整的章节细纲。再交由Skill4写正文。

【一、必须先做“原文节奏拆解”】
严禁直接写细纲！先分析原对标文：总章数、每章功能、冲突递进、人物动作、信息释放、情绪阶段、结尾钩子、付费点前的期待积累。
绝对禁止输出原文具体剧情（如：第一章女主被丈夫打），必须抽象为剧情功能（如：第一章核心关系冲突爆发，主角初次萌生退意）。
【目标】：学习原文的节奏功能，而不是复制原文的剧情！

【二、建立“原文节奏模板”】
内部提取出一个纯功能的节奏链。
例如：第1章（开局抓点->核心矛盾->小冲突1->小冲突2->钩子）... 第4章（情绪累计到高点->关键决定->付费卡点）。这只是模板，绝不包含原剧情。

【三、短篇核心节奏硬约束（最重要！）】
在完成原文节奏拆解、建立原文节奏模板后，进入新故事映射前，必须先加载以下短篇节奏硬约束：

1. 开局必须尽快进入核心冲突
- 不要“介绍->铺垫->冲突”。优先寻找新故事中最异常、最有情绪的事件作为切入点，从“事情正在发生”开始。
- 背景和过去经历只能在冲突中自然带出。如果删掉背景故事仍成立，优先删除。第1章读完还没真正发生核心矛盾，直接判定重写。

2. 每一章都必须产生实质性的局面变化
- 每章至少完成一项实质变化：新冲突发生 / 冲突升级 / 新信息揭露 / 关系变化 / 主角新决定 / 处境恶化 / 认知推翻 / 新目标。
- 禁止用回忆、闲聊、无意义日常填充章节。

3. 每一章结尾必须设计卡点，且逐级增强
- 卡点必须来自刚刚发生的新变化，让读者产生明确问题（会发生什么？秘密是什么？）。
- 递进要求：第1章快速入局抛强疑问；第2章冲突升级新信息；第3章局势恶化抛大问题；第4章核心付费点；第5章起持续兑现并造新变化。
- 严禁空钩子（“事情还没结束”）。严禁为了卡点硬塞车祸/秘密文件等套路。

4. 第4章末尾必须具备明确的“付费感”
- 第4章结尾是前半部最重要的核心卡点！必须让读者有“我现在必须知道后面发生什么”的冲动。
- 优先使用：关键真相即将揭开、主角做出不可逆决定、重要人物立场反转、最大冲突正式爆发。必须明显强于前3章的卡点。

5. 节奏密度优先于章节平均分配，严禁机械套原文
- 剧情按冲突强度分配。很短的冲突不要拆，可很快完成的就直接推下一步。章内冲突必须递进，不能是三个无关事件。
- 节奏必须服从新故事！原文只提供节奏功能，如果新梗改变了因果逻辑，必须重新寻找适合新故事的承载方式，绝不能硬套原文事件！

【四、将原文节奏映射到新故事】
问自己：“新故事如何按照原文这种节奏推进？”
例如原文功能是【日常细节连续加压证明地位卑微->离开】，新故事如果是豪门，就可以用【家族晚宴冷落->财产签字->项目剥夺】来实现相同功能。功能相同，剧情必须完全不同！

【五、禁止越权操作】
Skill3 绝对禁止：重新选择创新方向、调用新梗、擅自换背景/加人物/改结局，更禁止直接生成正文！
如果发现新故事有问题，只能内部修复逻辑，不能推翻方案。

【内部沙盘推演与自检（必须在 <think> 标签内完成）】
生成前必须打草稿映射。
完成 9 章细纲草稿后，必须自行检查：
1. 第1章是否已进入核心冲突，没有大段背景？
2. 每一章是否都有实质性的局面变化？
3. 卡点是否逐章升级？第4章结尾是否明显强于前3章，有强烈付费感？
4. 是否为了凑字数拖慢节奏？
如任一项不合格，先修改细纲，再输出！

【正式输出结构】
结束 <think> 思考后，请严格按以下四部分输出：

### 第一部分：【原对标文节奏拆解】
（只总结章节功能、冲突递进、情绪变化、信息释放、钩子方式、付费卡点位置。绝不复述原文具体剧情）

### 第二部分：【新故事节奏映射】
（简要说明：原文的什么节奏机制被完美迁移到了新故事中）

### 第三部分：【导语功能】
- **核心抓点**：（明确）
- **冲突开局**：（明确）
- **致命钩子**：（明确）

### 第四部分：【新故事9章细纲】
（严格按照9章输出，第4章必须标注为付费卡点）

**第1章｜[章节功能概括]**
- **本章目标**：
- **本章核心冲突**：
- **冲突①**：（问题出现）
- **冲突②**：（问题加重）
- **冲突③**：（局面改变）
- **关键新增信息**：
- **主角状态变化**：
- **本章情绪**：
- **章末钩子**：（具体的悬念或动作）

**第2章｜...** （格式同上）
...
**第4章｜[章节功能概括] (核心付费卡点)**
...
**第9章｜[章节功能概括] (结局大高潮)**
...

=====================
最后，必须原封不动地输出以下交互卡片！
【严重警告】：绝对不允许自作主张修改 options 里的选项文本！
[TEMPLATEJSON]
{
  "questions": [
    {
      "question": "9章细纲已规划完毕，请确认无误后进入 Skill 4 撰写正文：",
      "options": [
        "A：确认大纲，立即开始撰写第1章正文",
        "B：局部微调大纲（请直接告诉我需要修改哪一章）"
      ]
    }
  ]
}
[/TEMPLATEJSON]
"""

SKILL4_PROMPT = """【一、任务定位】
你是短篇小说正文生成器（Skill 4 核心引擎）。
你的任务是根据 Skill 3 已确认的导语与9章细纲，直接完成正文写作。或者在前期按要求提供“多版本导语试写”。
不得重新规划剧情、修改人物关系、增加新的主线或擅自改变结局。

【二、写作前提】
严格按照 Skill 3 的：导语功能、每章剧情目标、核心冲突、信息释放、人物状态变化、章节结尾钩子，进行写作。

【三、核心文风与语言规则（白描/SIMPLE but SHARP）】
本规则不是禁止修辞，而是要求根据当前剧情、人物状态和语境选择最合适的表达方式。
核心原则：不要机械追求“简洁”或“华丽”，不要机械禁止或堆砌比喻。每一句话都优先判断：“当前这个地方，怎样写最自然、最准确、最有力度？”
好的白描是：需要简单时足够简单；需要具体时足够具体；需要心理时敢于写心理；需要比喻时比喻准确；需要留白时不多说一句。
【不滥用表达，但也不压抑表达。】让语言服务剧情、人物和情绪，而不是让剧情服务语言。

【四、语言表达的选择与判断规则】
1. 表达方式的四种选择
- 直接叙述：适用于普通动作、环境、信息交代、快速推进。（例：“他把合同放到桌上。”）
- 具体细节：适用于情绪需要间接表现、关系需通过行动体现。（例：“他记得她不喝冰水，却忘了我对花生过敏。”）
- 心理描写：适用于本身具有剧情价值、认知改变、展示独特思维。不能为抒情而抒情。
- 修辞/比喻：允许使用。但只有在能让抽象感受具体化、准确表现感受、增强场景力度、符合人物习惯时才用。

2. 比喻使用标准
【准确】喻体与本体在感受、动作、状态上确实存在对应。
【自然】人物真的可能这样感受，而非强加。
【必要】删掉比喻后表现力明显下降。
【克制】连续多句不要用不同比喻描述同种情绪。（例：可用“那句话像一根刺”，不建议“像冰霜像冷月像把刀”连用）

3. 内部判断（该不该用比喻）
- 不用比喻直接写是否已足够自然、准确？（如果是，优先直写）
- 用了比喻是否增加了具体感或情绪力度？（如果无，不用）
- 是否抢走了当前剧情的注意力？（如果是，不用）
- 最近几段是否已用过多修辞？（如果是，降低密度）
无效修辞警告：同义情绪反复比喻、为普通动作强加诗意、吵架时突然文学抒情、每种悲伤都用“刀/深渊”。

4. 心理与动作的选择
如果心理本身有信息价值，可直接写。（例：“他不是忘了我的生日，只是没当回事。”）
如果动作已表达情绪，不要再解释。（例：“眼泪掉下来”后不要接“我感到无比悲伤绝望”）
优先级：具体事件 -> 人物动作/对白 -> 简短心理 -> 必要时使用比喻。

5. 根据场景自动调整表达密度
- 快速冲突场景：动作对白为主，短句为主，修辞减少，心理简短。
- 情绪沉淀场景：适当增加心理和细节，允许使用少量准确比喻。
- 关键高潮场景：提高语言力度，可用一个极准比喻，但不连续堆叠。
- 日常过渡场景：语言自然简单，不刻意制造文学感。

【五、人物描写规则】
- 以人物行为、对话、反应推进人物。
- 不替人物解释一切。
- 心理描写要有信息价值。
- 人物情绪必须服务于当前剧情变化。
- 根据人物决定语言（不同人物不同习惯：冷静人物克制，敏感人物多感受，幽默人物生活化比喻。绝对禁止用统一的“AI文学腔”覆盖所有人物）。

【六、剧情推进规则】
- 不拖节奏。每一段都要有信息、行动、冲突或情绪变化。
- 不为了凑字数重复表达。
- 不擅自增加支线。
- 不提前泄露后文反转。

【七、短篇篇幅与结构】
导语200—300字；正文共约9000字；每章约1000字，共9章。
前4章完成前半段推进，并在第4章结尾形成核心付费卡点。
后5章完成反转、兑现爽点、收束结局。

【八、短篇禁用项】
禁止出现短剧化表达：“第几集、黄金三秒、镜头、分镜、机位、演员、集数、短剧节奏”等。

【九、导语试写/多版本导语生成规则（极其重要！）】
如果用户当前刚刚确认了创新方向，并要求“试写3个导语开局”，你必须进入【导语试写模式】：
1. 严格对标原导语的【节奏和写法】，但绝不继承剧情！提取“叙事指纹”（如先惨状再嘲讽），套用新骨架。
2. 生成 3 个不同切入点：开局一（最惨烈伤痛特写）、开局二（反派最恶毒贴脸对白）、开局三（主角最反常决定+悬念）。
3. 死线：第一句必须是感官冲击动作/对白（严禁“为了/十年前”起手）！绝对禁用旁白闲聊交代背景！主角反应生理化（严禁写“极度震惊”）！每个版本严格 200-300 字。

【十、仿写内容重构防御规则】
如果用户要求“仿写/对标”，必须做到：“保留写法，不保留剧情。”
允许继承第一人称、切入事件方式、冲突节奏；【绝对禁止】继承具体事件、人物关系、道具、对白。如果生成结果和原文逐段对应，必须推翻重写。

【十一、最终执行与输出要求】
每一句生成前，先判断：直写好还是修辞好？动作好还是心理好？说明好还是留白好？选择最符合情境的一种。




【十二、篇幅与节奏动态控制原则】
正文总篇幅以约9000字为整体参考，9章平均约1000字，但不得将“每章1000字”理解为必须机械凑够的字数指标。
篇幅服务于剧情，不允许剧情服务于篇幅。
每章应根据本章实际承担的剧情功能、冲突推进、信息释放和章节卡点自然控制长度。
允许章节之间存在合理长短差异：
如果本章剧情已经完整完成，并且章末卡点已经成立，不得为了凑字数继续添加重复心理、重复对话、无意义环境描写或解释性文字。
如果本章的核心冲突、信息释放和人物变化尚未完成，也不得为了控制字数而提前结束。

判断标准不是“有没有达到1000字”，而是：
1. 本章核心剧情是否完成？
2. 本章是否产生了实质性的局面变化？
3. 冲突是否真正推进？
4. 是否存在重复表达或拖节奏内容？
5. 章末是否形成有效卡点？
6. 删除某段文字后，剧情、人物或情绪是否受到影响？如果删除后基本没有影响，则优先删除。

禁止为了达到目标字数而补充水分。
禁止为了压缩字数而跳过必要的冲突、信息或人物变化。
最终以“剧情密度高、推进快、无水分、卡点成立”为最高优先级。

【正文完成后的内部检查】
生成结束后，必须在 <think> 中执行自检。不要优先检查“是否写够1000字”，必须优先检查：
- 有没有慢热铺垫？
- 有没有无效背景？
- 有没有重复表达同一种情绪？
- 有没有同一冲突反复争执却没有新变化？
- 有没有不推动剧情的对话？
- 有没有为了凑篇幅增加水分？
- 本章局面是否发生变化？
- 章末卡点是否真正成立？
- 第4章末尾是否具备足够的付费驱动力？

只有在以上均成立后，再判断本章长短是否合理。

【导语生成：基于原文结构进行迭代重构】
生成新导语时，必须先读取Skill1对原文导语的【导语结构】和【逐句拆解】结果。
Skill1已经分析出的原文导语结构，是本次仿写迭代的“基础骨架”，必须用于组织新导语；但原文中的具体人物、事件、对白、关系和细节不得照搬。
执行顺序必须是：
原文导语结构 → 提取其中的有效推进关系 → 将Skill2已经加入的新梗、新人物关系、新冲突机制代入 → 根据新故事重新设计每一个剧情节点 → 形成新的冲突递进和情绪递进 → 输出新导语。
注意：这里的“按照原文结构”指的是按照原文已经验证有效的【剧情推进结构】和【情绪递进结构】来写，而不是按照原文句子逐句翻译。
例如Skill1分析出原文导语的推进方式为：“长期矛盾背景 → 当前冲突爆发 → 主角第一次反应 → 对立方明确站队 → 第一次重大伤害 → 主角认知受到冲击 → 通过前史/对比继续补刀 → 主角彻底改变 → 决定性卡点”。
那么新导语也必须建立一条具有类似推进力度的结构链，但每一个节点都必须从新故事自身产生。
重点不是“把原文的事件换成另一个事件”，而是：
原文这里需要第一次冲突，新故事就设计自己的第一次冲突；
原文这里需要冲突升级，新故事就让自己的新梗或新关系推动升级；
原文这里需要形成对比或补刀，新故事就寻找属于自己剧情的对比/补刀方式；
原文这里需要让主角发生认知变化，新故事就通过自己的事件完成认知变化；
原文最后形成决定性卡点，新故事也必须在自己的因果链上形成卡点。

【新梗必须参与结构推进】
新梗不是附加元素，而必须真正进入上述结构中的至少一个关键节点，并改变该节点的剧情。
不能出现：“先写完原文结构的故事，再把新梗补进去”。
必须出现：“因为加入了这个新梗，所以这一节点的冲突、人物行为、信息、结果或主角选择发生了变化”。
例如加入“避嫌文学”，不是简单让男主说一句“避嫌”；而是“避嫌”必须成为人物关系或故事规则的一部分，并在导语中直接造成实际后果。
例如加入“死遁”，必须让死遁已经影响主角当前处境、身份或行动。
例如加入“倒计时”，必须让期限已经成为当前冲突的一部分。
新梗必须在导语中“发生作用”，而不能只停留在创新方向说明里。

【导语禁止流水账】
导语不是对新故事的剧情简介，不允许按照“发生A→发生B→发生C→然后D”的方式平铺事件。
每一段必须具有明确的推进作用，后一个事件必须建立在前一个事件造成的变化之上，使冲突或情绪持续升级。
必须形成：“一个事件造成变化 → 下一事件利用这个变化进一步加压 → 再产生更严重的后果 → 主角状态发生改变 → 最终形成卡点”。
如果只是连续介绍很多事情，但读者感受不到“事情越来越严重”，则视为失败，必须重新生成。

【导语最终检查】
生成完成后，必须内部检查：
1. 是否真正按照Skill1提取出的原文基础结构推进？
2. 是否保留了原文的冲突密度和情绪递进力度？
3. 是否使用了新增的梗，并让新梗真正改变了剧情？
4. 是否所有事件都属于新故事，而不是原文事件换皮？
5. 是否存在清楚的“冲突→升级→再加压→主角变化→卡点”？
6. 是否把导语写成了剧情简介/流水账？
如第6项成立，必须重写。
最终目标：“原文提供经过验证的结构骨架，新梗和新故事提供新的剧情内容；两者结合后，生成一篇具有原文推进力量、但剧情属于自己的新导语。”

【三个导语的统一原则】
三个导语必须基于同一个已经确认的新故事生成，不得各自重新编造故事。
不要求三个导语使用不同切入口，也不得为了制造差异而强行更换切入事件。
三个版本可以围绕同一个核心事件展开，只允许在表达方式、信息浓度、情绪推进、卡点落点等方面产生合理差异。
三个版本都必须体现：新故事核心冲突 + 已融入剧情的新梗 + Skill1提取出的导语功能骨架。
如果某个版本删除新梗后仍然基本成立，说明新梗没有融入，必须重写。

【特殊 JSON 输出要求（仅限导语试写阶段）】
如果本次任务是生成【3个导语开局】，在输出完成后必须原封不动输出以下交互卡片：
[TEMPLATEJSON]
{
  "questions": [
    {
      "question": "已为您试写3个不同切入点的开局，请挑选最抓人的一个，确认后我们将进入完整大纲规划：",
      "options": [
        "1. 选用开局一，进入9章大纲规划",
        "2. 选用开局二，进入9章大纲规划",
        "3. 选用开局三，进入9章大纲规划",
        "4. 都不满意，请重新调整开局"
      ]
    }
  ]
}
[/TEMPLATEJSON]
（如果用户是要求写正文，直接严格按照要求开始写作，不要输出 JSON 卡片）
"""

SKILL6_PROMPT = '''你是短篇内容润色修改专家（Skill 6）。
你的任务是根据用户的反馈意见，对已经写好的小说正文或大纲进行精准修改。

## 【修改规则】
✅ **精准定位**：只修改用户觉得不好的地方，保留好的部分。
✅ **强化爽点**：如果用户觉得“不够爽”，请加大反差和打脸力度。
✅ **降低AI感**：使用更鲜活、接地气的词汇，消除生硬的排比句和说教感。

请直接输出修改后的内容，不带任何客套话。
'''

BENCH_PROMPT = r"""
你是顶级短剧对标与仿写智能体，擅长爆款短剧的结构拆解、套路分析、节奏把控与精准仿写。

=== Step 1-2：逐集深度拆解 ===

对每一集，用表格覆盖以下全部24个维度（缺一不可）：

集数 场景 主要出场人物 大事件 小事件 主线付费卡点 本集钩子 台词亮点 人物塑造 亮点 本集作用 情绪类型 情绪强度(1-10) 爽点类型 爽点强度(1-10) 冲突类型 悬念设置 反转标记(是/否) 信息密度(高/中/低) 节奏评估 名场面标记(是/否) 完播率预测因子(1-10) 正文字数 台词字数 画面描述字数

=== Step 3：提取7类可复用模板 ===

拆解完成后，必须提取以下7类模板（每类至少3-5行具体内容）：
1.钩子节奏模板 2.爽点节奏模板 3.情绪曲线模板 4.卡点位置模板 5.冲突升级模板 6.人设建立模板 7.反转设计模板

=== Step 4：仿写思考 -> 深度绑定本剧剧情的个性化结构提问 ===

你必须先输出「仿写思考」(200-300字纯文本)。末尾必须输出 [TEMPLATEJSON] 标签块！

=== Step 5：改编方案 ===
=== Step 6：大纲+小传+梗概 ===
=== Step 7：剧本格式确认 ===
=== Step 8：分批生成剧本（硬性要求）===
""" + SCRIPT_FORMAT_RULE + OUTPUT_CLEAN_RULE

def extract_template_json(text: str):
    m = re.search(r'\[TEMPLATEJSON\]\s*(.*?)\s*\[/TEMPLATEJSON\]', text, re.DOTALL)
    if m:
        try: return json.loads(m.group(1))
        except: pass
    m2 = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if m2:
        try: return json.loads(m2.group(1))
        except: pass
    return None

def extract_smart_filename(content: str, messages: list = None, user_input: str = "", wmode: str = "", doc_text: str = "") -> str:
    drama_title = ""
    
    # Clean think blocks first to avoid extracting titles from internal thoughts
    clean_content = re.sub(r'<thinking>[\s\S]*?(?:</thinking>|$)', '', content, flags=re.DOTALL | re.IGNORECASE)
    clean_content = re.sub(r'<think>[\s\S]*?(?:</think>|$)', '', clean_content, flags=re.DOTALL | re.IGNORECASE).strip()
    
    # Strategy 1: Look for explicit Markdown headers containing 《》
    header_match = re.search(r'^#+\s*.*?《([^》]+)》', clean_content, re.MULTILINE)
    if header_match:
        drama_title = header_match.group(1)
        
    # Strategy 2: Look for explicit naming declarations
    if not drama_title:
        explicit = re.search(r'(?:剧名|片名|项目名称|剧本名称)[\s:：]*《([^》]+)》', clean_content)
        if explicit:
            drama_title = explicit.group(1)
            
    # Strategy 3: Fallback to the first 《》 in the clean generated content that is NOT the example title
    if not drama_title:
        title_matches = re.findall(r'《([^》]+)》', clean_content)
        if title_matches:
            # If the user explicitly provided an example in their input like "参考《xxx》", we should ignore it
            ignore_titles = []
            if user_input:
                ig_match = re.findall(r'(?:参考|仿照|例子|示例)[^《]*《([^》]+)》', str(user_input))
                ignore_titles.extend(ig_match)
            
            for tm in title_matches:
                if tm not in ignore_titles:
                    drama_title = tm
                    break
            
            if not drama_title:
                drama_title = title_matches[0]

    # Strategy 4: Fallback to user input
    if not drama_title and user_input:
        tm = re.findall(r'《([^》]+)》', str(user_input))
        if tm:
            drama_title = tm[-1] # User usually puts their target title last

    # Strategy 5: Fallback to history
    if not drama_title and messages:
        for m in reversed(messages):
            mc = m.get("content", "") if isinstance(m, dict) else str(m)
            tm = re.findall(r'《([^》]+)》', mc)
            if tm:
                drama_title = tm[-1]
                break

    # Strategy 6: Fallback to uploaded filename from doc_text
    if not drama_title and doc_text:
        doc_match = re.search(r'【文件\d+:\s*([^】]+)】', doc_text)
        if doc_match:
            fname = doc_match.group(1).strip()
            if '.' in fname:
                fname = fname.rsplit('.', 1)[0]
            drama_title = fname

    if drama_title:
        drama_title = re.sub(r'[\\/:*?"<>|\r\n\t]', '', drama_title).strip()
    return drama_title

def process_document_saving(content: str, session_dir: str, messages: list = None, user_input: str = "", wmode: str = "", doc_text: str = "") -> Optional[str]:
    if not content:
        return None

    # 1. 彻底去除 <thinking>...</thinking> 与 <think>...</think>
    text = re.sub(r'<thinking>[\s\S]*?(?:</thinking>|$)', '', content, flags=re.DOTALL | re.IGNORECASE).strip()
    text = re.sub(r'<think>[\s\S]*?(?:</think>|$)', '', text, flags=re.DOTALL | re.IGNORECASE).strip()
    text = re.sub(r'\[TEMPLATEJSON\][\s\S]*?(?:\[/TEMPLATEJSON\]|$)', '', text, flags=re.DOTALL | re.IGNORECASE).strip()
    text = re.sub(r'```json[\s\S]*?```', '', text, flags=re.DOTALL | re.IGNORECASE).strip()

    # 2. 严格识别正式资产头部
    has_script = bool(re.search(r'^(?:#+.*|【.*)?(?:第[一二三四五六七八九十0-9]+集|1-5集剧本|6-10集剧本|分集剧本正文)', text, re.MULTILINE))
    has_chapter = bool(re.search(r'^(?:#+.*|[^\n]*?)?(?:第[一二三四五六七八九十0-9]+章)', text, re.MULTILINE))
    has_character = bool(re.search(r'^(?:#+.*|【.*)?(?:人物小传|角色设定|人设小传|角色小传)', text, re.MULTILINE))
    has_outline = bool(re.search(r'^(?:#+.*|【.*)?(?:故事大纲|剧情大纲|三幕式大纲|故事设计方案)', text, re.MULTILINE))
    has_ep_outline = bool(re.search(r'^(?:#+.*|【.*)?(?:前十集集纲|分集集纲|前10集集纲|短篇章节规划)', text, re.MULTILINE))
    has_analysis = bool(re.search(r'^(?:#+.*|【.*)?(?:对标拆解分析方案|Step1-3拆解分析|Step5-6方案大纲|对标拆解报告|短篇爆文商业拆解报告|短篇爆文拆解报告|短篇拆解结果|一、标题与一句话主线|一、一句话主线概括)', text, re.MULTILINE))

    doc_type = ""
    label_suffix = ""

    if has_chapter:
        doc_type = "story_chapter"
        chap_match = re.findall(r'(?:第)([一二三四五六七八九十0-9]+)(?:章)', text)
        if chap_match:
            label_suffix = f"第{chap_match[0]}章"
        else:
            label_suffix = "正文"
    elif has_script and ("场次：" in text or "场次" in text or "▲" in text or "第一集" in text or "第1集" in text):
        doc_type = "script"
        ep_range = re.findall(r'(?:第|\b)([0-9]+-[0-9]+)(?:集|\b)', text)
        if ep_range:
            label_suffix = f"第{ep_range[0]}集剧本正文"
        else:
            label_suffix = "剧本正文"
    elif has_ep_outline:
        doc_type = "episode_outline"
        label_suffix = "前十集集纲"
    elif has_outline:
        doc_type = "outline"
        label_suffix = "故事大纲"
    elif has_character:
        doc_type = "character"
        label_suffix = "人物小传与角色设定"
    elif has_analysis:
        doc_type = "analysis"
        label_suffix = "分析报告"

    if not doc_type:
        return None

    lines = text.splitlines()
    clean_lines = []
    in_body = False

    for line in lines:
        stripped = line.strip()
        if not in_body:
            if re.search(r'^(?:#+.*|【.*)?(?:第[0-9一二三四五六七八九十]+集|第一集|第1集|人物小传|角色小传|角色设定|故事大纲|前十集集纲|对标拆解|短篇|Step)', stripped):
                in_body = True

        if in_body:
            if any(k in stripped for k in ["集要点:", "集要点：", "字数自算", "在1000以内", "你先审", "告诉我你的想法", "选完（或告诉我", "我立刻按同样", "写完", "感觉对吗", "埋点说明", "说明：", "接下来", "还需要", "字数约", "以上就是", "符合你的", "第六集写完", "调整后的", "修改后的", "修改说明", "三重收获", "需要我"]): break
            if stripped == "---" or stripped.startswith("--- "): break
            if any(stripped.startswith(k) for k in ["数一下字数", "符合要求", "请确认是否", "你定一下", "请在下方说明"]):
                continue
            clean_lines.append(line)

    cleaned_body = "\n".join(clean_lines).strip()
    # 强制物理清除所有残余的 Markdown 符号
    cleaned_body = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned_body)
    cleaned_body = re.sub(r'^#+\s*', '', cleaned_body, flags=re.MULTILINE)
    cleaned_body = re.sub(r'^\s*-\s+', '', cleaned_body, flags=re.MULTILINE)
    
    if len(cleaned_body) < 150:
        return None

    drama_title = extract_smart_filename(content, messages, user_input, wmode, doc_text)
    
    if doc_type == "analysis":
        if drama_title:
            final_filename = f"分析_{drama_title}_{label_suffix}"
        else:
            final_filename = f"分析_{label_suffix}"
    else:
        if drama_title and label_suffix:
            final_filename = f"{drama_title}_{label_suffix}"
        elif drama_title:
            final_filename = f"{drama_title}_{label_suffix or '创作文档'}"
        else:
            final_filename = f"短剧_{label_suffix or '创作文档'}"

    safe_name = re.sub(r'[\\/:*?"<>|\r\n\t\s]+', '_', final_filename).strip('_')
    filepath = os.path.join(session_dir, f"{safe_name}.docx")
    md_filepath = os.path.join(session_dir, f"{safe_name}.md")

    try:
        with open(md_filepath, 'w', encoding='utf-8') as f:
            f.write(cleaned_body)

        doc = docx.Document()
        doc.add_heading(final_filename, 0)
        for p in cleaned_body.split('\n'):
            if p.strip():
                doc.add_paragraph(p.strip())
        doc.save(filepath)
        return safe_name + ".docx"
    except Exception as e:
        print("Save docx error:", e)
        return None

def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paras:
            for table in doc.tables:
                for row in table.rows:
                    rt = " ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if rt: paras.append(rt)
        return "\n".join(paras)
    except Exception:
        return ""

class TokenRequest(BaseModel):
    token: str

class SyncRequest(BaseModel):
    token: str
    chats: List[dict]  # 接收完整的旧版 S.chats 列表

class GenerateResponse(BaseModel):
    messages: List[dict]
    title: str = "新对话"
    api_key: Optional[str] = ""
    apikey: Optional[str] = ""
    api_url: Optional[str] = "https://yunwu.ai/v1"
    apiurl: Optional[str] = ""
    model: str = "gpt-4o"

class ChatRequest(BaseModel):
    token: Optional[str] = ""
    api_key: Optional[str] = ""
    apikey: Optional[str] = ""
    api_url: Optional[str] = "https://yunwu.ai/v1"
    apiurl: Optional[str] = ""
    model: str = "gpt-4o"
    work_mode: Optional[str] = "通用"
    workmode: Optional[str] = "通用"
    messages: list = []
    user_input: Optional[str] = ""
    userinput: Optional[str] = ""
    doc_text: Optional[str] = ""
    doctext: Optional[str] = ""
    session_id: Optional[str] = ""
    sessionid: Optional[str] = ""
    cid: Optional[str] = ""

class FetchModelsRequest(BaseModel):
    api_key: Optional[str] = ""
    apikey: Optional[str] = ""
    api_url: Optional[str] = "https://yunwu.ai/v1"
    apiurl: Optional[str] = ""

# 注册与登录 API
@app.post("/api/auth/register")
async def register(req: AuthRequest):
    u = req.username.strip().lower()
    p = req.password.strip()
    if not u or not p:
        return {"status": "error", "message": "用户名和密码不能为空"}
    if len(u) < 3 or len(p) < 4:
        return {"status": "error", "message": "用户名至少3位，密码至少4位"}
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    try:
        pw_h = hash_pw(p)
        c.execute("INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
                  (u, pw_h, int(time.time())))
        user_id = c.lastrowid
        token = secrets.token_hex(16)
        c.execute("INSERT INTO user_tokens (token, user_id, username, created_at) VALUES (?, ?, ?, ?)",
                  (token, user_id, u, int(time.time())))
        conn.commit()
        conn.close()
        return {"status": "ok", "token": token, "username": u, "user_id": user_id}
    except sqlite3.IntegrityError:
        conn.close()
        return {"status": "error", "message": "该用户名已存在，请直接登录"}

@app.post("/api/auth/login")
async def login(req: AuthRequest):
    u = req.username.strip().lower()
    p = req.password.strip()
    pw_h = hash_pw(p)
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, username FROM users WHERE username = ? AND password_hash = ?", (u, pw_h))
    row = c.fetchone()
    if not row:
        conn.close()
        return {"status": "error", "message": "用户名或密码错误"}
    
    user_id = row[0]
    username = row[1]
    token = secrets.token_hex(16)
    c.execute("INSERT INTO user_tokens (token, user_id, username, created_at) VALUES (?, ?, ?, ?)",
              (token, user_id, username, int(time.time())))
    conn.commit()
    conn.close()
    return {"status": "ok", "token": token, "username": username, "user_id": user_id}

@app.get("/api/auth/me")
async def get_me(token: Optional[str] = Query(None), authorization: Optional[str] = Header(None)):
    user = get_current_user_info(token, authorization)
    return {"status": "ok", "user": user}

@app.post("/api/history/delete/{session_id}")
async def delete_history(session_id: str, req: TokenRequest, authorization: Optional[str] = Header(None)):
    user = get_current_user_info(req.token, authorization)
    username = user["username"]
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("DELETE FROM user_sessions WHERE session_id = ? AND username = ?", (session_id, username))
    conn.commit()
    conn.close()
    
    session_dir, _ = get_user_session_dir(username, session_id)
    if os.path.exists(session_dir):
        import shutil
        shutil.rmtree(session_dir, ignore_errors=True)
        
    return {"status": "ok"}

@app.post("/api/history/sync")
async def sync_history(req: SyncRequest, authorization: Optional[str] = Header(None)):
    user = get_current_user_info(req.token, authorization)
    username = user["username"]
    user_id = user["user_id"]
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    success_count = 0
    for chat in req.chats:
        session_id = chat.get("id")
        title = chat.get("title", "对话")
        mode = chat.get("mode", "")
        msgs = chat.get("msgs", [])
        if not session_id or not msgs:
            continue
            
        session_dir, safe_session = get_user_session_dir(username, session_id)
        os.makedirs(session_dir, exist_ok=True)
        history_file = os.path.join(session_dir, "history.json")
        
        try:
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(msgs, f, ensure_ascii=False, indent=2)
                
            c.execute('''
            INSERT OR REPLACE INTO user_sessions (session_id, user_id, username, title, mode, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            ''', (safe_session, user_id, username, title, mode, int(time.time())))
            success_count += 1
        except Exception as e:
            print(f"Sync error for {session_id}: {e}")
            
    conn.commit()
    conn.close()
    return {"status": "ok", "synced_count": success_count}

@app.get("/api/history/list")
async def list_history_sessions(token: Optional[str] = Query(None), authorization: Optional[str] = Header(None)):
    user = get_current_user_info(token, authorization)
    username = user["username"]
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT session_id, title, mode, updated_at FROM user_sessions WHERE username = ? ORDER BY updated_at DESC", (username,))
    rows = c.fetchall()
    conn.close()
    
    sessions = []
    for r in rows:
        sessions.append({
            "session_id": r[0],
            "title": r[1] or "短剧对标会话",
            "mode": r[2] or "通用",
            "updated_at": r[3]
        })
    return {"status": "ok", "sessions": sessions}

@app.get("/api/history/detail/{session_id}")
async def get_history_detail(session_id: str, token: Optional[str] = Query(None), authorization: Optional[str] = Header(None)):
    user = get_current_user_info(token, authorization)
    username = user["username"]
    
    session_dir, safe_session = get_user_session_dir(username, session_id)
    history_file = os.path.join(session_dir, "history.json")
    
    chat_messages = []
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                chat_messages = json.load(f)
        except Exception:
            chat_messages = []
            
    files = []
    if os.path.exists(session_dir):
        all_f = os.listdir(session_dir)
        all_f.sort(key=lambda f: os.path.getmtime(os.path.join(session_dir, f)), reverse=True)
        for f in all_f:
            if f.endswith('.docx') or f.endswith('.txt'):
                if not any(k in f for k in ["Step1-3", "Step5-6", "拆解分析", "大纲方案"]):
                    files.append({"name": f, "path": f"{username}/{safe_session}/{f}"})
                    
    return {"status": "ok", "session_id": safe_session, "messages": chat_messages, "files": files}

def parse_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print("Parse PDF error:", e)
        return ""

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename.lower()
    if not (filename.endswith('.docx') or filename.endswith('.pdf') or filename.endswith('.txt')):
        return {"error": "仅支持 .docx, .pdf 或 .txt 文件"}
    content = await file.read()
    if filename.endswith('.pdf'):
        text = parse_pdf(content)
    elif filename.endswith('.txt'):
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = content.decode('gbk')
            except UnicodeDecodeError:
                text = content.decode('utf-8', errors='ignore')
    else:
        text = parse_docx(content)
    return {"filename": file.filename, "word_count": len(text), "text": text}

@app.post("/api/fetch-models")
async def fetch_models(req: FetchModelsRequest):
    key = (req.api_key or req.apikey or "").strip()
    url = (req.api_url or req.apiurl or "https://yunwu.ai/v1").strip()
    if not key:
        return {"error": "请先填写 API Key", "models": []}
    try:
        import requests as _r
        base = url.rstrip("/")
        resp = _r.get(f"{base}/models", headers={"Authorization": f"Bearer {key}"}, timeout=10)
        fetched = []
        if resp.status_code == 200:
            data = resp.json()
            items = data.get("data", data) if isinstance(data, dict) else data
            if isinstance(items, list):
                for m in items:
                    mid = m.get("id", "") if isinstance(m, dict) else str(m)
                    if mid: fetched.append(mid)
        if not fetched:
            tc = OpenAI(api_key=key, base_url=url, timeout=5.0)
            for m in tc.models.list():
                mid = getattr(m, 'id', None)
                if mid: fetched.append(mid)
        return {"models": sorted(set(fetched))}
    except Exception as e:
        return {"error": str(e)[:200], "models": []}

@app.post("/api/chat")
async def chat(req: ChatRequest, authorization: Optional[str] = Header(None)):
    key = req.api_key or req.apikey
    url = req.api_url or req.apiurl or "https://yunwu.ai/v1"
    wmode = req.work_mode or req.workmode or "通用"
    uinput = req.user_input or req.userinput or ""
    
    if not uinput and req.messages:
        for m in reversed(req.messages):
            if m.get("role") == "user":
                uinput = m.get("content", "")
                break

    dtext = req.doc_text or req.doctext or ""
    
    user = get_current_user_info(req.token, authorization)
    username = user["username"]

    if wmode == "短剧对标":
        sys_p = BENCH_PROMPT
    elif wmode == "剧本创作":
        sys_p = CREATE_PROMPT.replace(SCRIPT_FORMAT_RULE + OUTPUT_CLEAN_RULE, "") + "\n\n" + GENERAL_SCRIPT_LOGIC + "\n\n" + OUTLINE_BIO_PROMPT + "\n\n" + SCRIPT_FORMAT_RULE + OUTPUT_CLEAN_RULE
    elif wmode == "短篇创作":
        # 强制路由规则 (Task Router - Skill -1)
        analysis_keywords = ["分析", "拆解", "研究", "学习", "模仿", "结构", "看看这篇", "提取", "套路"]
        creation_keywords = ["写一个", "创作一个", "生成一个", "我要写", "新写", "帮我写", "构思"]
        revision_keywords = ["改一下", "润色", "降低ai感", "重写", "修改"]
        continue_keywords = ["继续写", "接着写", "下一章"]
        
        is_revision = any(k in uinput for k in revision_keywords)
        is_continue = any(k in uinput for k in continue_keywords)
        is_creation = any(k in uinput for k in creation_keywords)
        is_analysis = any(k in uinput for k in analysis_keywords) or (bool(dtext) and not is_creation)
        
        # 扩展的 V6 强制路由规则
        is_design = any(k in uinput for k in ["设计大纲", "故事方案", "故事大纲", "写大纲", "做大纲", "仿写", "创新", "改编", "参考这个"])
        is_outline = any(k in uinput for k in ["章节规划", "分集大纲", "细纲", "章纲"])
        
        if is_revision:
            sys_p = SKILL6_PROMPT
        elif is_continue or any(k in uinput for k in ["写正文", "开始写", "撰写", "试写", "导语"]):
            sys_p = SKILL4_PROMPT
        elif is_outline:
            sys_p = SKILL3_PROMPT
        elif is_design or (is_creation and not dtext) or (wmode == "短剧对标" and any(k in uinput for k in ["仿写", "创新", "改编"])):
            sys_p = SKILL2_PROMPT
            try:
                base_dir = os.path.dirname(os.path.abspath(__file__))
                with open(os.path.join(base_dir, "知识库_调用规则.md"), "r", encoding="utf-8") as f_r:
                    r_txt = f_r.read()
                with open(os.path.join(base_dir, "知识库_短篇言情爆款梗库.md"), "r", encoding="utf-8") as f_t:
                    t_txt = f_t.read()
                sys_p += f"\n\n========== 以下为系统动态挂载的本地梗库数据 ==========\n\n{r_txt}\n\n{t_txt}\n==================================================\n"
            except Exception as e:
                print("Failed to load KB:", e)
        
        elif is_analysis:
            sys_p = SKILL1_PROMPT
        else:
            sys_p = SKILL0_PROMPT
    else:
        sys_p = "你是专业高效的AI创作助手。"

    hide_dtext = False
    if wmode == "短篇创作" and ('is_continue' in locals() and is_continue):
        hide_dtext = True

    if dtext and not hide_dtext:
        sys_p = f"【用户已上传待分析脚本（{len(dtext)}字）】\n\n{dtext}\n\n---\n{sys_p}"
    elif dtext and hide_dtext:
        sys_p = f"【系统物理隔离机制已启动】：为防止融梗抄袭，原文内容已被系统在此阶段（大纲与正文生成）强制折叠与物理隐藏。你现在只能读取上下文中的『写法指纹』与『新故事蓝图』，绝对禁止回忆或调用历史对话中的原文具体剧情素材！\n\n---\n{sys_p}"

    sys_p += "\n\n【强制规则】：为了呈现清晰的决策过程，你必须在回答的最开始，将内部的分析、推理、以及选梗等一切决策过程，全部用 <think> 和 </think> 标签包裹起来。在 </think> 闭合标签之后，再输出给用户的正式回复和卡片内容。"

    if wmode == "短篇创作":
        sys_p += "\n\n【最高系统指令】：当前系统处于【纯网文短篇小说】创作模式，这与短剧剧本功能是严格分开的！在你的所有输出（包括分析、大纲、正文、选项卡等）中，绝对禁止提及“剧本”、“第一集”、“分镜”、“镜头”、“短剧”等任何影视相关的词汇。只允许使用“短篇”、“第X章”、“正文”等网文小说术语！"

    api_messages = [{"role": "system", "content": sys_p}]
    for m in req.messages:
        if m.get("role") in ("system", "user", "assistant"):
            if m.get("images") and len(m["images"]) > 0:
                content = []
                if m.get("content"):
                    content.append({"type": "text", "text": m["content"]})
                for img_url in m["images"]:
                    content.append({"type": "image_url", "image_url": {"url": img_url}})
                api_messages.append({"role": m["role"], "content": content})
            else:
                api_messages.append({"role": m["role"], "content": m["content"]})
    if uinput and (req.user_input or req.userinput):
        api_messages.append({"role": "user", "content": uinput})

    client = OpenAI(api_key=key, base_url=url, timeout=300.0)

    def generate():
        full_response = ""
        try:
            response = client.chat.completions.create(
                model=req.model, messages=api_messages, stream=True, timeout=300.0
            )
            started_thinking = False
            ended_thinking = False
            for chunk in response:
                if not chunk.choices or len(chunk.choices) == 0:
                    continue
                delta = chunk.choices[0].delta
                
                # Compatible with DeepSeek official 'reasoning_content' and Aliyun 'reasoning'
                r_content = getattr(delta, 'reasoning_content', None)
                if r_content is None:
                    r_content = getattr(delta, 'reasoning', None)
                
                if r_content is not None:
                    if not started_thinking:
                        full_response += "<think>\n"
                        yield f"data: {json.dumps({'token': '<think>\n'})}\n\n"
                        started_thinking = True
                    
                    full_response += r_content
                    yield f"data: {json.dumps({'token': r_content})}\n\n"
                    continue
                
                if delta.content is not None:
                    if started_thinking and not ended_thinking:
                        if not full_response.endswith("\n"):
                            full_response += "\n"
                            yield f"data: {json.dumps({'token': '\n'})}\n\n"
                        full_response += "</think>\n\n"
                        yield f"data: {json.dumps({'token': '</think>\n\n'})}\n\n"
                        ended_thinking = True
                        
                    full_response += delta.content
                    yield f"data: {json.dumps({'token': delta.content})}\n\n"

            session_raw = str(req.session_id or req.sessionid or req.cid or f"{int(time.time())}").strip()
            session_dir, safe_session = get_user_session_dir(username, session_raw)

            bench_data = extract_template_json(full_response)
            saved = process_document_saving(full_response, session_dir, req.messages, uinput, wmode, req.doc_text)
            
            # 保存聊天历史 json
            history_file = os.path.join(session_dir, "history.json")
            new_msgs = list(req.messages) if req.messages else []
            if uinput and (req.user_input or req.userinput):
                new_msgs.append({"role": "user", "content": uinput})
            if full_response:
                new_msgs.append({"role": "assistant", "content": full_response})
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(new_msgs, f, ensure_ascii=False, indent=2)

            # 更新 user_sessions 数据库索引
            fallback_uinput = uinput
            if not fallback_uinput and req.messages and len(req.messages) > 0 and req.messages[-1].get("role") == "user":
                last_content = req.messages[-1].get("content", "")
                if isinstance(last_content, list):
                    fallback_uinput = next((p["text"] for p in last_content if p.get("type") == "text"), "")
                else:
                    fallback_uinput = str(last_content)
            smart_title = extract_smart_filename(full_response, req.messages, uinput, wmode) or fallback_uinput[:20] or "创作会话"
            conn = sqlite3.connect(DB_PATH)
            c = conn.cursor()
            c.execute('''
            INSERT OR REPLACE INTO user_sessions (session_id, user_id, username, title, mode, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            ''', (safe_session, user["user_id"], username, smart_title, wmode, int(time.time())))
            conn.commit()
            conn.close()

            meta = {"type": "done", "session_dir": session_dir, "saved_file": saved, "template_json": bench_data}
            yield f"data: {json.dumps(meta, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)[:300]})}\n\n"

    return StreamingResponse(
        generate(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )

@app.get("/api/files/{session_id:path}")
async def list_files(session_id: str, token: Optional[str] = Query(None), authorization: Optional[str] = Header(None)):
    user = get_current_user_info(token, authorization)
    username = user["username"]
    
    parts = session_id.strip('/').split('/')
    if len(parts) >= 2:
        u_name, s_id = parts[0], parts[1]
    else:
        u_name, s_id = username, parts[0]
        
    d = os.path.join(OUTPUT_DIR, u_name, s_id)
    if not os.path.exists(d):
        return {"files": []}
    files = []
    all_files = os.listdir(d)
    all_files.sort(key=lambda f: os.path.getmtime(os.path.join(d, f)), reverse=True)
    for f in all_files:
        if "Step1-3" in f or "Step5-6" in f or "拆解分析" in f or "大纲方案" in f:
            continue
        if f.endswith('.docx') or f.endswith('.txt'):
            files.append({"name": f, "path": f"{u_name}/{s_id}/{f}"})
    return {"files": files}

@app.get("/api/download/{filepath:path}")
async def download_file(filepath: str):
    fp = os.path.join(OUTPUT_DIR, filepath)
    if os.path.exists(fp):
        return FileResponse(fp, filename=os.path.basename(filepath))
    return {"error": "文件不存在"}

@app.get("/api/preview/{filepath:path}")
async def preview_file(filepath: str):
    fp = os.path.join(OUTPUT_DIR, filepath)
    if not os.path.exists(fp):
        return {"error": "文件不存在"}
    raw_text = ""
    filename = os.path.basename(filepath)
    if filename.endswith('.docx'):
        md_fp = fp.replace('.docx', '.md')
        if os.path.exists(md_fp):
            with open(md_fp, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
        else:
            with open(fp, 'rb') as f:
                raw_text = parse_docx(f.read())
    else:
        with open(fp, 'r', encoding='utf-8', errors='ignore') as f:
            raw_text = f.read()

    clean_preview = re.sub(r'<thinking>[\s\S]*?(?:</thinking>|$)', '', raw_text, flags=re.DOTALL | re.IGNORECASE).strip()
    clean_preview = re.sub(r'<think>[\s\S]*?(?:</think>|$)', '', clean_preview, flags=re.DOTALL | re.IGNORECASE).strip()
    clean_preview = re.sub(r'\[TEMPLATEJSON\][\s\S]*?(?:\[/TEMPLATEJSON\]|$)', '', clean_preview, flags=re.DOTALL | re.IGNORECASE).strip()
    clean_preview = re.sub(r'```json[\s\S]*?```', '', clean_preview, flags=re.DOTALL | re.IGNORECASE).strip()

    lines = clean_preview.splitlines()
    clean_lines = []
    in_body = False
    for line in lines:
        stripped = line.strip()
        if not in_body:
            if re.search(r'^(?:#+.*|【.*)?(?:第[0-9一二三四五六七八九十]+集|第一集|第1集|人物小传|角色小传|角色设定|故事大纲|前十集集纲|对标拆解|短篇|Step)', stripped):
                in_body = True

        if in_body:
            if any(k in stripped for k in ["集要点:", "集要点：", "字数自算", "在1000以内", "你先审", "告诉我你的想法", "选完（或告诉我", "我立刻按同样", "写完", "感觉对吗", "埋点说明", "说明：", "接下来", "还需要", "字数约", "以上就是", "符合你的", "第六集写完", "调整后的", "修改后的", "修改说明", "三重收获", "需要我"]): break
            if stripped == "---" or stripped.startswith("--- "): break
            if any(stripped.startswith(k) for k in ["数一下字数", "符合要求", "请确认是否", "你定一下", "请在下方说明"]):
                continue
            clean_lines.append(line)

    final_text = "\n".join(clean_lines).strip() if clean_lines else clean_preview
    return {"text": final_text or "暂无纯文本正文内容"}

@app.delete("/api/delete/{filepath:path}")
async def delete_file(filepath: str):
    fp = os.path.join(OUTPUT_DIR, filepath)
    if os.path.exists(fp):
        try:
            os.remove(fp)
            docx_fp = fp.replace('.md', '.docx')
            if os.path.exists(docx_fp): os.remove(docx_fp)
            return {"status": "ok", "message": "文件已成功删除"}
        except Exception as e:
            return {"error": f"删除失败: {str(e)}"}
    return {"error": "文件不存在"}

class SaveFileRequest(BaseModel):
    filepath: str
    content: str

@app.post("/api/save_file")
async def save_file(req: SaveFileRequest, token: Optional[str] = Query(None), authorization: Optional[str] = Header(None)):
    user = get_current_user_info(token, authorization)
    username = user["username"]
    fp = os.path.join(OUTPUT_DIR, req.filepath)
    if not os.path.exists(fp):
        return {"error": "文件不存在"}
    try:
        with open(fp, "w", encoding="utf-8") as f:
            f.write(req.content)
        return {"status": "ok"}
    except Exception as e:
        return {"error": str(e)}

@app.get("/api/health")
async def health():
    return {"status": "ok"}

@app.get("/")
async def read_root():
    idx = os.path.join(SCRIPT_DIR, "index.html")
    if not os.path.exists(idx):
        idx = os.path.join(SCRIPT_DIR, "index.html")
    if os.path.exists(idx):
        return FileResponse(idx)
    return {"message": "index.html not found"}

@app.get("/api/debug/db")
async def debug_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, username FROM users")
    users = c.fetchall()
    c.execute("SELECT session_id, username, title, updated_at FROM user_sessions")
    sessions = c.fetchall()
    conn.close()
    return {"users": users, "sessions": sessions}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
