import streamlit as st
from openai import OpenAI
import docx, os, time, json, re

st.set_page_config(page_title="创作工坊", page_icon="🎬", layout="wide")

if not hasattr(st, "rerun"):
    st.rerun = getattr(st, "experimental_rerun", st.rerun)

THEMES = {
    "🍑 蜜桃乌龙": {"bg":"#fef9f4","secondarybg":"#fdf0e5","text":"#5c4a3d","primary":"#e8937a","cardbg":"#ffffff","border":"#f5d5c6","shadow":"0 4px 20px rgba(232,147,122,0.12)","radius":"12px"},
    "🫧 气泡苏打": {"bg":"#f6fafe","secondarybg":"#eaf4fd","text":"#3d5064","primary":"#7eb8da","cardbg":"#ffffff","border":"#d0e7f5","shadow":"0 4px 20px rgba(126,184,218,0.10)","radius":"16px"},
    "🍋 柠檬跳跳糖": {"bg":"#fffef5","secondarybg":"#fff9db","text":"#5a5530","primary":"#f0c040","cardbg":"#ffffff","border":"#f5e8a0","shadow":"0 4px 20px rgba(240,192,64,0.10)","radius":"10px"},
    "🧸 小熊软糖": {"bg":"#fbf7f4","secondarybg":"#f5ede6","text":"#5c4638","primary":"#c4956a","cardbg":"#ffffff","border":"#e8d5c4","shadow":"0 4px 20px rgba(196,149,106,0.10)","radius":"14px"},
    "🖋 墨纸禅意": {"bg":"#fafaf8","secondarybg":"#f2f2ee","text":"#2c2c2c","primary":"#3a3a3a","cardbg":"#ffffff","border":"#e0e0dc","shadow":"0 2px 12px rgba(0,0,0,0.04)","radius":"6px"},
    "🌙 月光灰调": {"bg":"#f5f5f7","secondarybg":"#eaeaef","text":"#2d2d3a","primary":"#5b5b8a","cardbg":"#ffffff","border":"#dcdce5","shadow":"0 2px 16px rgba(91,91,138,0.06)","radius":"8px"},
    "🪨 暖石灰": {"bg":"#f9f6f3","secondarybg":"#f0ece8","text":"#4a4540","primary":"#8b7e6e","card_bg":"#ffffff","border":"#e2ddd6","shadow":"0 2px 12px rgba(139,126,110,0.06)","radius":"8px"},
    "🌿 薄荷暗色": {"bg":"#1e2328","secondarybg":"#161a1e","text":"#d0d6d8","primary":"#68b893","cardbg":"#262c32","border":"#353d45","shadow":"0 4px 20px rgba(0,0,0,0.25)","radius":"10px"},
    "🫐 蓝莓暗夜": {"bg":"#1a1d2e","secondarybg":"#141726","text":"#c8cddb","primary":"#7c8ce0","cardbg":"#232740","border":"#353b58","shadow":"0 4px 20px rgba(0,0,0,0.30)","radius":"10px"},
    "🖤 曜石极简": {"bg":"#111111","secondarybg":"#0a0a0a","text":"#cccccc","primary":"#eeeeee","cardbg":"#1a1a1a","border":"#333333","shadow":"0 4px 20px rgba(0,0,0,0.40)","radius":"4px"},
}

HISTORYFILEPATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "chats_history.json")

def loadhistorychatsfromfile():
    if os.path.exists(HISTORYFILEPATH):
        try:
            with open(HISTORYFILEPATH, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    return data
        except Exception:
            return []
    return []

def savehistorychatstofile():
    try:
        with open(HISTORYFILEPATH, "w", encoding="utf-8") as f:
            json.dump(st.session_state.historychats, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

if "initdone" not in st.session_state:
    st.session_state.historychats = loadhistorychatsfromfile()
    st.session_state.currentchatid = f"chat_{int(time.time())}"
    st.session_state.workmode = None
    st.session_state.messages = [{"role":"system","content":"你是专业高效的AI创作助手。"}]
    st.session_state.sessiondir = None
    st.session_state.benchstep4_template = None
    st.session_state.benchstep4_confirmed = False
    st.session_state.benchstep7_confirmed = False
    st.session_state.benchbatch_index = 0
    st.session_state.benchtotal_batches = 0
    st.session_state.currenttheme_name = "🍑 蜜桃乌龙"
    st.session_state.fetchedmodel_list = ["deepseek-chat","deepseek-reasoner","gpt-4o","gpt-4o-mini"]
    st.session_state.activepreview_content = None
    st.session_state.activepreview_title = ""
    st.session_state.isgenerating = False
    st.session_state.stoprequested = False
    st.session_state.editingidx = None
    st.session_state.autocontinue = False
    st.session_state.initdone = True

def archivecurrentchat():
    msgs = st.session_state.messages
    user_msgs = [m["content"] for m in msgs if m["role"] == "user"]
    if user_msgs:
        firstq = user_msgs[0]
        firstq = re.sub(r'[#*`\[\]【】\s]', '', firstq)
        title = firstq[:14] if firstq else "短剧对话"
        found = False
        for c in st.session_state.historychats:
            if isinstance(c, dict) and c.get("id") == st.session_state.currentchatid:
                c["messages"] = list(msgs)
                c["workmode"] = st.session_state.workmode
                c["title"] = title
                c["sessiondir"] = st.session_state.sessiondir
                found = True
                break
        if not found:
            st.session_state.historychats.insert(0, {
                "id": st.session_state.currentchatid,
                "title": title,
                "messages": list(msgs),
                "workmode": st.session_state.workmode,
                "sessiondir": st.session_state.sessiondir
            })
        savehistorychatstofile()

def startnewchat():
    archivecurrentchat()
    st.session_state.currentchatid = f"chat_{int(time.time())}"
    st.session_state.messages = [{"role":"system","content":"你是专业高效的AI创作助手。"}]
    st.session_state.workmode = None
    st.session_state.sessiondir = None
    st.session_state.benchstep4_template = None
    st.session_state.benchstep4_confirmed = False
    st.session_state.benchstep7_confirmed = False
    savehistorychatstofile()

def loadchatbyid(chatid):
    archivecurrentchat()
    for c in st.session_state.historychats:
        if isinstance(c, dict) and c.get("id") == chatid:
            st.session_state.currentchatid = c["id"]
            st.session_state.messages = list(c.get("messages", []))
            st.session_state.workmode = c.get("workmode", None)
            st.session_state.sessiondir = c.get("sessiondir", None)
            st.session_state.benchstep4_template = None
            st.session_state.benchstep4_confirmed = False
            st.session_state.benchstep7_confirmed = False
            break
    savehistorychatstofile()

def deletechatbyid(chatid):
    st.session_state.historychats = [c for c in st.session_state.historychats if isinstance(c, dict) and c.get("id") != chatid]
    savehistorychatstofile()
    if st.session_state.currentchatid == chatid:
        startnewchat()

def createsessiondir():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    outdir = os.path.join(script_dir, "output")
    d = os.path.join(outdir, f"session_{int(time.time())}")
    os.makedirs(d, exist_ok=True)
    return d

def extracttemplatejson(text):
    if not text: return None
    m = re.search(r'\[TEMPLATEJSON\]\s*(.*?)\s*\[/TEMPLATEJSON\]', text, re.DOTALL)
    if m:
        try: return json.loads(m.group(1))
        except: return None
    return None

def clean_tags(text):
    return re.sub(r'\[TEMPLATEJSON\].*?\[/TEMPLATEJSON\]', '', text, flags=re.DOTALL)

def hidejson(text):
    return re.sub(r'\[TEMPLATEJSON\].*?\[/TEMPLATEJSON\]', '\n\n📋 *(模板数据已解析)*\n\n', text, flags=re.DOTALL)

def autosave(content, session_dir, label=None):
    try:
        clean = clean_tags(content)
        os.makedirs(session_dir, exist_ok=True)
        if label:
            safe_label = re.sub(r'[\\/:*?"<> ]', '', label)
            fname = f"{safe_label}.md"
        else:
            ts = int(time.time())
            fname = f"output_{ts}.md"
        with open(os.path.join(session_dir, fname), "w", encoding="utf-8") as f:
            f.write(clean)
        doc = docx.Document()
        doc.add_heading("创作报告", 0)
        for p in clean.split('\n'):
            if p.strip(): doc.add_paragraph(p.strip())
        doc.save(os.path.join(session_dir, fname.replace('.md', '.docx')))
        return (session_dir, fname)
    except Exception as e:
        return (None, f"Error:{e}")

def splitthinking(text):
    if not text: return None, text
    thinkmatch = re.search(r'<(?:thinking|thought)>(.*?)</(?:thinking|thought)>', text, re.DOTALL)
    if thinkmatch:
        thinking = thinkmatch.group(1).strip()
        body = text[:thinkmatch.start()] + text[thinkmatch.end():]
        return thinking, body.strip()
    return None, text

OUTPUTCLEANRULE = """
【思考与正文隔离铁律】：
请先在 <thinking> ... </thinking> 标签内进行编剧思考与拉片拆解。
在 <thinking> 标签外输出最终正文。正文禁止带有 ** 错乱符号。
"""

CREATEPROMPT = r"""
你是顶级竖屏短剧编剧智能体。必须严格按以下方法论从零到一完成短剧剧本创作。

核心理论：
情绪ABC短剧变形：情绪 = 视觉符号(B)  冲击动作(A)  秒级节奏(T)
压弹簧理论：每一集只有压弹簧(积蓄期待)或放弹簧(释放爽感)两种状态。
爆款核心逻辑：隐藏身份 + 场景冲突 + 打脸反转 + 追悔莫及。

创作前确认参数：总集数、单集字数(+-10%)、输出格式、台词占比(默认65%)。

模式1-头脑风暴：核心人设5要素-核心场景-情绪主轴-终极反转-付费卡点。
模式2-大纲规划：三幕式(1-10建立期待-11-60积累压抑-61-80爆发释放)。五秒节拍法。
模式3-人物小传：六维人设容器(标签/观念/视觉符号/防御机制/核心目标/最大矛盾)。反派动机必须真实。
模式4-剧本优化：冰山四维度。每场5-8句台词，有来有回层层加码。

四类剧本格式：真人竖屏/AI仿真人/AI漫剧/传统横屏。
分批生成(每轮5集)：轮次回顾-逐集生成-字数校验-台词占比校验-继续。
13条铁律：3秒定生死/每15秒转折/观念靠看不靠说/反派嚣张3倍/主角隐忍/冲击动作有声音/反应镜头不能省/卡点是开始/同一符号反复用/情绪要纯/字数+-10%/台词60%-70%/每轮5集。
""" + OUTPUTCLEANRULE

BENCHPROMPT = r"""
你是顶级短剧对标与仿写智能体。你必须严格按照以下8步工作流执行，任何步骤不得跳过或简并。

=== Step 1-2：逐集深度拆解 ===

对每一集，用表格覆盖以下全部24个维度（缺一不可）：

集数 场景 主要出场人物 大事件 小事件 主线付费卡点 本集钩子 台词亮点 人物塑造 亮点 本集作用 情绪类型 情绪强度(1-10) 爽点类型 爽点强度(1-10) 冲突类型 悬念设置 反转标记(是/否) 信息密度(高/中/低) 节奏评估 名场面标记(是/否) 完播率预测因子(1-10) 正文字数 台词字数 画面描述字数

情绪类型：爽/虐/甜/悬/怒/悲/喜/恐
爽点类型：逆袭/打脸/甜宠/复仇/虐渣/揭秘/救赎/共鸣
冲突类型：人物/利益/情感/身份/阶级/时间
完播率预测公式：情绪强度0.35 + 爽点强度0.25 + 钩子强度(有钩子=3/强悬念=5/付费卡点=5)0.25 + 信息密度(高=2/中=1/低=0)0.15

=== Step 3：提取7类可复用模板 ===

拆解完成后，必须提取以下7类模板（每类至少3-5行具体内容）：

1.钩子节奏模板：逐集列出钩子类型+钩子内容+是否付费卡点，总结钩子覆盖率
2.爽点节奏模板：逐集爽点类型+强度，标注蓄力期与释放期的集数边界
3.情绪曲线模板：逐集情绪类型+强度，标注压弹簧/放弹簧的节奏规律
4.卡点位置模板：逐集卡点位置，计算卡点间隔规律
5.冲突升级模板：分阶段描述冲突类型变化（每阶段2-3集），标注升级方式
6.人设建立模板：主角+反派+2个关键配角，每人分阶段(初始-触发-挣扎-转变-终点)
7.反转设计模板：逐集列出反转内容+类型+前置铺垫（标出铺垫集数）

末尾输出核心商业模型：标题公式 + 核心公式 + 弹簧节奏公式

在拆解正文末尾，嵌入[TEMPLATEJSON]块，包含prototypepatterns(每类模板用name和summary)、referencetitle、genre、episodecount、corehook。不用带step4_questions。

=== Step 4：仿写思考 -> 提问（严禁输出JSON问题列表）===

你必须先输出「仿写思考」再提问。思考+问题都用纯文本，禁止用JSON包裹问题。

「仿写思考」(200-300字)必须包含：
本剧最独特的3个记忆锚点——具体的名场面/视觉符号/台词，不能抽象
核心人物关系模型——谁和谁是情感引擎？
弹簧节奏的独特性——压/放的比例和强度
如果只保留一样东西到仿写中，你选什么？为什么？

然后基于思考，在正文中提出6个开放式问题。每个问题必须引用本剧的具体元素。示例：
"《白骨精》里骨妖替祁娘而活+抚养砖儿柳儿+反杀武二的核心公式，你打算怎么处理？贴身保留还是大改？"
禁止："是否保留核心设定？"

问题涵盖：仿写方向、新设定（谁替谁活）、篇幅节奏、情绪定位、差异化锚点、受众平台。

全文末尾只输出一行：[TEMPLATEJSON]{"step":"step4_ready"}[/TEMPLATEJSON]

然后立即停止，等待用户在面板中填写方案。

=== Step 5：改编方案 ===

用户确认后，生成包含6个章节的改编方案。

=== Step 6：大纲+小传+梗概 ===

包含剧本大纲(500-800字)+主角小传+分集梗概(每集3-5句)+字数预算表。

=== Step 6.5：字数预算 ===

快节奏集：预算下限-10%至-5%。中节奏集：预算中点。慢节奏集：预算上限+5%至+10%。
估算公式：场景头约50字 + 画面块约40字 + 对话轮约50字 + 钩子约30字。
末尾输出[TEMPLATEJSON]{"step":"step7_format"}[/TEMPLATEJSON]

=== Step 7：剧本格式确认 ===

输出默认格式供参考，等待用户确认。

=== Step 8：分批生成剧本（硬性要求）===

禁止一次性输出全部集数。每轮3-5集。
每轮流程：轮次提示 -> Step 8.3回顾 -> 逐集生成 -> Step 8.5字数校验 -> 轮次汇报 -> 引导下一轮。
Step 8.3回顾(非首轮)：逐项检查剧情推进/钩子衔接/人物状态/冲突进度/情绪曲线/节奏密度/字数合规。
Step 8.5字数校验：统计画面描述+台词+钩子的纯正文字数。偏差超过+-10%立即修正。
每轮末输出[TEMPLATEJSON]{"step":"batchcomplete","batchindex":N,"total_batches":M}[/TEMPLATEJSON]

=== 输出控制 ===

单次输出不超过6000字，超出标[请说继续获取下一段]
拆解表格24行必须全部填满，禁止空行或"—"
JSON必须在[TEMPLATEJSON]和[/TEMPLATEJSON]之间
""" + OUTPUTCLEANRULE

# ---- CSS ----
theme = THEMES[st.session_state.currenttheme_name]
is_dark = theme["bg"].startswith("#1")
css = ":root{--dot:" + theme["primary"] + ";}"
css += "@keyframes td{0%,60%,100%{opacity:0.2;transform:translateY(0)}30%{opacity:1;transform:translateY(-6px)}}"
css += "@keyframes slideUp{from{opacity:0;transform:translateY(12px)}to{opacity:1;transform:translateY(0)}}"
css += ".stApp{background:" + theme["bg"] + ";color:" + theme["text"] + ";transition:background .5s ease,color .5s ease;}"
css += "section[data-testid=\"stSidebar\"]{background:" + theme["secondarybg"] + ";border-right:1px solid " + theme["border"] + ";}"
css += ".stButton>button{border-radius:" + theme["radius"] + ";font-weight:500;background:" + theme["cardbg"] + ";border:1px solid " + theme["border"] + ";box-shadow:" + theme["shadow"] + ";color:" + theme["text"] + ";transition:all .25s;}"
css += ".stButton>button:hover{transform:translateY(-2px);border-color:" + theme["primary"] + ";color:" + theme["primary"] + ";}"
css += ".stTextInput>div>div>input,.stSelectbox>div>div>div,.stTextArea textarea{background:" + theme["cardbg"] + ";color:" + theme["text"] + ";border-radius:8px;border:1px solid " + theme["border"] + ";}"
css += "div[data-testid=\"stChatMessage\"]{background:" + theme["cardbg"] + ";border-radius:14px;border:1px solid " + theme["border"] + ";box-shadow:" + theme["shadow"] + ";padding:16px;margin-bottom:12px;animation:slideUp .3s ease;}"
css += ".panel-box{background:" + theme["cardbg"] + ";border:2px solid " + theme["primary"] + ";border-radius:14px;padding:20px;margin:12px 0;box-shadow:" + theme["shadow"] + ";}"
css += ".panel-box h4{color:" + theme["primary"] + ";}"
errbg = "#3a1a1a" if is_dark else "#fff0f0"
errborder = "#ff6b6b" if is_dark else "#ffcccc"
css += ".error-box{background:" + errbg + ";border:1px solid " + errborder + ";border-radius:10px;padding:14px 18px;margin:4px 0;}"
css += ".error-box .err-title{color:#ff6b6b;font-weight:600;font-size:14px;margin-bottom:4px;}"
css += ".error-box .err-body{color:" + theme["text"] + ";font-size:13px;opacity:.8;margin-bottom:10px;}"
css += ".thinking-box{background:" + theme["secondarybg"] + ";border:1px dashed " + theme["border"] + ";border-radius:10px;padding:12px 16px;margin-bottom:12px;font-size:13px;opacity:.85;}"
css += ".thinking-box summary{color:" + theme["primary"] + ";font-weight:600;cursor:pointer;}"
st.markdown("<style>" + css + "</style>", unsafe_allow_html=True)

# ---- 侧边栏 ----
with st.sidebar:
    st.markdown("### 主题切换")
    themenames = list(THEMES.keys())
    currentidx = themenames.index(st.session_state.currenttheme_name)
    selected = st.selectbox("选主题", themenames, index=currentidx, key="theme_select")
    if selected != st.session_state.currenttheme_name:
        st.session_state.currenttheme_name = selected
        st.rerun()

    st.divider()
    st.markdown("### 历史对话")
    cnew, chist = st.columns([1.2, 1.8])
    with cnew:
        if st.button("+ 新建对话", use_container_width=True, key="btnnewchat"):
            startnewchat()
            st.rerun()
    with chist:
        historylist = st.session_state.historychats
        if historylist and isinstance(historylist, list):
            validchats = [h for h in historylist if isinstance(h, dict)]
            if validchats:
                labels = []
                curindex = 0
                for idxh, h in enumerate(validchats):
                    t = h.get("title", "历史对话")
                    iscur = (h.get("id") == st.session_state.currentchatid)
                    if iscur:
                        curindex = idxh
                    labels.append(f"{'👉 ' if iscur else '  '}{t}")
                if curindex >= len(labels):
                    curindex = 0
                try:
                    selectedlabel = st.selectbox("历史记录", labels, index=curindex, key="histselect", label_visibility="collapsed")
                except TypeError:
                    selectedlabel = st.selectbox("历史记录", labels, index=curindex, key="hist_select")
                selidx = labels.index(selectedlabel) if selectedlabel in labels else 0
                targeth = validchats[selidx]
                if targeth.get("id") != st.session_state.currentchatid:
                    loadchatbyid(targeth["id"])
                    st.rerun()

    st.divider()
    with st.expander("API设置", expanded=False):
        apikey = st.text_input("API Key", value="", type="password", key="apikeyinput")
        apiurl = st.text_input("API URL", value="https://yunwu.ai/v1", key="apiurlinput")
        st.divider()
        st.markdown("模型选择")
        if st.button("拉取可用模型", key="btn_fetch"):
            if apikey and apiurl:
                fetched = []
                errdetail = ""
                try:
                    import requests as _r
                    base = apiurl.rstrip("/")
                    resp = _r.get(f"{base}/models", headers={"Authorization": f"Bearer {apikey}"}, timeout=10)
                    if resp.status_code == 200:
                        data = resp.json()
                        items = data.get("data", data) if isinstance(data, dict) else data
                        if isinstance(items, list):
                            for m in items:
                                mid = m.get("id", "") if isinstance(m, dict) else str(m)
                                if mid: fetched.append(mid)
                    if not fetched:
                        tc = OpenAI(api_key=apikey, base_url=apiurl, timeout=5.0)
                        for m in tc.models.list():
                            mid = getattr(m, 'id', None)
                            if mid: fetched.append(mid)
                except Exception as e:
                    errdetail = str(e)[:150]
                if fetched:
                    st.session_state.fetchedmodel_list = sorted(set(fetched))
                    st.success(f"拉取成功：{len(fetched)} 个模型")
                else:
                    st.error(f"拉取失败，请检查API Key和URL是否正确。\n{errdetail}")
            else:
                st.warning("请先填写API Key和URL")
        selectedmodel = st.selectbox("选择模型", st.session_state.fetchedmodel_list, key="sel_model")
        custommodel = st.text_input("或手动输入模型ID", value="", key="custom_model")
        modelname = custommodel.strip() if custommodel.strip() else selectedmodel

    st.divider()
    st.header("参考脚本上传")
    uploadedfile = st.file_uploader("上传Word脚本(.docx)", type=["docx"], key="file_upload")
    doctext = ""
    if uploadedfile is not None:
        try:
            doc = docx.Document(uploadedfile)
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paras:
                for table in doc.tables:
                    for row in table.rows:
                        rt = " ".join([c.text.strip() for c in row.cells if c.text.strip()])
                        if rt: paras.append(rt)
            doctext = "\n".join(paras)
            if doctext:
                st.success(f"已加载:{uploadedfile.name}({len(doctext)}字)")
            else:
                st.warning("未提取到文字")
        except Exception as e:
            st.error(f"读取失败:{e}")

    st.divider()
    st.header("项目文件")
    sd = st.session_state.sessiondir
    if sd and os.path.exists(sd):
        afs = []
        for root, dirs, files in os.walk(sd):
            for f in files:
                if f.endswith(('.md','.docx','.txt')):
                    afs.append((os.path.relpath(os.path.join(root,f),sd), os.path.join(root,f)))
        if afs:
            with st.expander("生成文件", expanded=True):
                for rel, full in afs:
                    c1,c2 = st.columns([5,1])
                    with c1:
                        if st.button(f"{rel}", key=f"tree{rel}", use_container_width=True):
                            try:
                                if rel.endswith('.docx'):
                                    st.session_state.activepreview_content = f"[Word]{full}"
                                else:
                                    with open(full,"r",encoding="utf-8") as fh:
                                        st.session_state.activepreview_content = fh.read()
                                st.session_state.activepreview_title = rel
                            except Exception as err:
                                st.session_state.activepreview_content = f"读取失败:{err}"
                                st.session_state.activepreview_title = rel
                            st.rerun()
                    with c2:
                        try:
                            if rel.endswith('.docx'):
                                st.download_button("down",open(full,"rb"),file_name=rel,key=f"dl_{rel}")
                            else:
                                with open(full,"r",encoding="utf-8") as fh:
                                    st.download_button("down",fh.read(),file_name=rel,key=f"dl_{rel}")
                        except:
                            pass

# ---- 主界面 ----
st.markdown("""
<div style='text-align:center;padding:20px 0 10px 0;'>
    <h3 style='font-weight:600;font-size:18px;margin:0;'>Hello，今天想创作些什么？</h3>
    <p style='font-size:12px;opacity:0.6;margin-top:4px;'>输入你的需求，开始创作吧！</p>
</div>
""", unsafe_allow_html=True)

if st.session_state.isgenerating:
    if st.button("停止生成", key="btn_stop"):
        st.session_state.stoprequested = True
        st.rerun()

if st.session_state.activepreview_content:
    with st.expander(f"预览:{st.session_state.activepreview_title}", expanded=True):
        st.markdown(st.session_state.activepreview_content)

# ---- Step 4 面板 ----
userchoiceinput = None

if st.session_state.workmode == "短剧对标":
    benchmsgs = st.session_state.messages

    if not st.session_state.benchstep4_confirmed:
        for m in reversed(benchmsgs):
            if m["role"] == "assistant" and not m.get("is_error"):
                data = extracttemplatejson(m["content"])
                if data and ("step4questions" in data or "step4_questions" in data or data.get("step") in ["step4ready","step4_ready"]):
                    st.session_state.benchstep4_template = data
                    break

    step4 = st.session_state.benchstep4_template
    if step4 and not st.session_state.benchstep4_confirmed:
        with st.container():
            st.markdown("---")
            st.markdown(f"### 「{step4.get('reference_title','未知')}」仿写方案填写")
            st.caption("AI已在聊天中提出6个问题，参考后在此自由描述你的仿写方案。")
            cf = st.text_area(
                "你的仿写方案（自然语言描述）：",
                placeholder="例：贴身仿写《猎户真千金》，保留男扮女装+替身+憨拳三件套。新设定：退役女特种兵替死去的豪门私生女...",
                key="s4_freeform",
                height=200
            )
            if st.button("确认方案，生成Step5&6", key="sub_s4", use_container_width=True):
                st.session_state.benchstep4_confirmed = True
                userchoiceinput = f"以下是我对仿写方案的描述，请基于此执行Step5和Step6：\n\n{cf.strip()}" if cf.strip() else "贴身仿写，保留原剧核心公式。请执行Step5和Step6。"
                st.rerun()
            st.markdown("---")

# ---- 聊天记录 ----
msgs = st.session_state.messages
for idx, message in enumerate(msgs):
    if message["role"] == "system":
        continue

    modetag = st.session_state.workmode or "general"
    msgkeybase = f"{modetag}{idx}"

    with st.chat_message(message["role"]):
        if message["role"] == "user":
            displaytext = message["content"]
            if st.session_state.editingidx == idx:
                newtext = st.text_area("编辑消息", value=displaytext, key=f"edit{msgkeybase}", height=100)
                c1, c2 = st.columns([1, 1])
                with c1:
                    if st.button("确认修改", key=f"confirm{msgkeybase}"):
                        msgs[idx]["content"] = newtext
                        del msgs[idx + 1:]
                        st.session_state.editingidx = None
                        st.rerun()
                with c2:
                    if st.button("取消", key=f"cancel{msgkeybase}"):
                        st.session_state.editingidx = None
                        st.rerun()
            else:
                st.markdown(displaytext)
                if not st.session_state.isgenerating and st.session_state.editingidx is None:
                    if st.button("edit", key=f"editbtn{msgkeybase}", help="编辑此消息"):
                        st.session_state.editingidx = idx
                        st.rerun()
            continue

        displaytext = hidejson(message["content"])
        CONTINUEMARKER = '[请说继续获取下一段]'
        hascontinue = CONTINUEMARKER in message["content"]
        displaytext = displaytext.replace(CONTINUEMARKER, '')

        thinking, body = splitthinking(displaytext)

        if message.get("is_error"):
            st.markdown(
                "<div class='error-box'><div class='err-title'>生成失败</div>"
                f"<div class='err-body'>{body}</div></div>",
                unsafe_allow_html=True
            )
        else:
            if thinking:
                with st.expander("思考过程", expanded=False):
                    st.markdown(thinking)
            st.markdown(body)

        islast = (message == msgs[-1])
        if message["role"] == "assistant" and islast and not st.session_state.isgenerating:
            if st.session_state.editingidx is None:
                if message.get("is_error"):
                    if st.button("重新生成", key=f"retryerr{msgkeybase}"):
                        msgs.pop()
                        st.rerun()
                else:
                    ca, cb = st.columns([1, 1])
                    with ca:
                        if st.button("重新生成", key=f"retry{msgkeybase}"):
                            msgs.pop()
                            st.rerun()
                    if hascontinue:
                        with cb:
                            if st.button("继续生成", key=f"cont{msgkeybase}"):
                                st.session_state.autocontinue = True
                                st.rerun()

# ---- Step 7 面板 ----
if st.session_state.workmode == "短剧对标" and st.session_state.benchstep4_confirmed and not st.session_state.benchstep7_confirmed:
    benchmsgs = st.session_state.messages
    for m in reversed(benchmsgs):
        if m["role"] == "assistant" and not m.get("is_error"):
            data = extracttemplatejson(m["content"])
            if data and data.get("step") == "step7_format":
                st.markdown("<div class='panel-box'><h4>剧本格式确认</h4></div>", unsafe_allow_html=True)
                with st.expander("默认格式参考（不填则用此格式）", expanded=True):
                    st.markdown("""
集数：第X集
场次：X-Y（换场景换场次）
场景：简洁（如"家里客厅""小区楼下便利店"）
人物：本集出场人物

▲画面描述：（简短，谁在哪里干什么，可见的动作和表情）

角色A：台词（口语化，15字以内为佳）

▲画面描述：...

角色B：台词

【本集钩子】

默认规则：每集1-2个场景 / 台词口语化15字以内 / 画面只写可见内容 / 不写比喻和内心独白 / 钩子必须在下一集开头回应
                    """)
                prompt_text = data.get("prompt", "")
                if prompt_text:
                    st.markdown(prompt_text)
                cf = st.text_area(
                    "描述你需要的剧本格式（留空=使用默认格式）：",
                    placeholder="例：每集3个场景、台词20字内、需标注情绪标签...",
                    key="s7_cf", height=100
                )
                if st.button("确认，生成剧本", key="sub_s7"):
                    st.session_state.benchstep7_confirmed = True
                    fm = cf.strip() if cf.strip() else "使用默认格式"
                    userchoiceinput = f"确认格式：{fm}。请执行Step8，生成第1轮(第1-5集)。"
                break

# ---- 分批继续面板 ----
if st.session_state.workmode == "短剧对标" and st.session_state.benchstep7_confirmed and st.session_state.benchbatch_index > 0:
    benchmsgs = st.session_state.messages
    bi = st.session_state.benchbatch_index
    tb = st.session_state.benchtotal_batches
    st.markdown(f"<div class='panel-box'><h4>第{bi}轮完成(共{tb}轮)</h4></div>", unsafe_allow_html=True)
    cc1, cc2 = st.columns([1, 1])
    with cc1:
        if bi < tb and st.button(f"继续第{bi+1}轮"):
            ns = bi * 5 + 1
            ne = min(ns + 4, tb * 5)
            userchoiceinput = f"执行Step8第{bi+1}轮(第{ns}-{ne}集)，先做回顾。"
            st.session_state.benchbatch_index = bi + 1
    with cc2:
        if st.button("导出全部"):
            ac = "\n\n---\n\n".join([
                hidejson(m["content"]) for m in benchmsgs
                if m["role"] == "assistant" and not m.get("is_error")
            ])
            sdd = st.session_state.sessiondir or createsessiondir()
            p, _ = autosave(ac, sdd, "完整导出")
            if p:
                st.success("已导出")

# ---- 底部按钮 + 输入框 ----
c1, c2, c3 = st.columns([1, 1, 8])

with c1:
    isbench = (st.session_state.workmode == "短剧对标")
    if st.button("对标" + (" V" if isbench else ""),
                 use_container_width=True, key="bt_bench"):
        st.session_state.workmode = None if isbench else "短剧对标"
        if isbench:
            st.session_state.benchstep4_confirmed = False
            st.session_state.benchstep7_confirmed = False
            st.session_state.benchbatch_index = 0
        st.rerun()

with c2:
    iscreate = (st.session_state.workmode == "剧本创作")
    if st.button("创作" + (" V" if iscreate else ""),
                 use_container_width=True, key="bt_create"):
        st.session_state.workmode = None if iscreate else "剧本创作"
        if not iscreate:
            st.session_state.benchstep4_confirmed = False
            st.session_state.benchstep7_confirmed = False
            st.session_state.benchbatch_index = 0
        st.rerun()

prompt = st.chat_input(f"在【{st.session_state.workmode or '通用'}】模式下输入指令...")

if st.session_state.get("autocontinue", False):
    st.session_state.autocontinue = False
    finalinput = "继续"
else:
    finalinput = prompt if prompt else (userchoiceinput if userchoiceinput else None)

if finalinput:
    if not apikey:
        st.error("请先填写API Key！")
    else:
        msgs = st.session_state.messages
        if st.session_state.workmode == "剧本创作":
            sysp = CREATEPROMPT
        elif st.session_state.workmode == "短剧对标":
            sysp = BENCHPROMPT
        else:
            sysp = "你是专业高效的AI创作助手。"

        if not msgs or msgs[0]["role"] != "system":
            msgs.insert(0, {"role":"system","content":sysp})
        else:
            msgs[0]["content"] = sysp

        if uploadedfile is not None and doctext and len(msgs) <= 2:
            msgs[0]["content"] = f"【用户已上传待分析脚本《{uploadedfile.name}》（{len(doctext)}字）】\n\n{doctext}\n\n---\n{msgs[0]['content']}"

        msgs.append({"role":"user","content":finalinput})
        archivecurrentchat()
        st.session_state.isgenerating = True
        st.session_state.stoprequested = False
        st.rerun()

# ---- API 调用 ----
if st.session_state.isgenerating and not st.session_state.stoprequested:
    if not apikey:
        st.error("请先填写API Key！")
        st.session_state.isgenerating = False
    else:
        msgs = st.session_state.messages
        apimessages = [{"role":m["role"],"content":m["content"]} for m in msgs]
        client = OpenAI(api_key=apikey, base_url=apiurl, timeout=300.0)

        with st.chat_message("assistant"):
            placeholder = st.empty()
            try:
                placeholder.markdown(
                    "<div style='color:#999;font-size:14px;padding:4px 0;'>"
                    "正在生成中<span style='display:inline-flex;align-items:center;gap:4px;margin-left:4px;'>"
                    "<span style='width:6px;height:6px;border-radius:50%;background:var(--dot);animation:td 1.4s infinite'></span>"
                    "<span style='width:6px;height:6px;border-radius:50%;background:var(--dot);animation:td 1.4s infinite .2s'></span>"
                    "<span style='width:6px;height:6px;border-radius:50%;background:var(--dot);animation:td 1.4s infinite .4s'></span>"
                    "</span></div>",
                    unsafe_allow_html=True
                )

                response = client.chat.completions.create(
                    model=modelname, messages=apimessages, stream=True, timeout=300.0
                )

                fullresponse = ""
                for chunk in response:
                    if st.session_state.stoprequested:
                        fullresponse += "\n\n[用户停止了生成]"
                        break
                    if not chunk.choices or len(chunk.choices) == 0:
                        continue
                    delta = chunk.choices[0].delta
                    if delta.content is not None:
                        fullresponse += delta.content
                        display = hidejson(fullresponse)
                        th, bd = splitthinking(display)
                        show = ""
                        if th:
                            show += f"<div class='thinking-box'><details><summary>思考过程</summary>{th}</details></div>\n\n"
                        show += bd
                        placeholder.markdown(show, unsafe_allow_html=True)

                display = hidejson(fullresponse)
                th, bd = splitthinking(display)
                show = ""
                if th:
                    show += f"<div class='thinking-box'><details open><summary>思考过程</summary>{th}</details></div>\n\n"
                show += bd
                placeholder.markdown(show, unsafe_allow_html=True)

                msgs.append({"role":"assistant","content":fullresponse})

                sdd = st.session_state.sessiondir
                if not sdd:
                    sdd = createsessiondir()
                    st.session_state.sessiondir = sdd

                savelabel = f"output_{int(time.time())}"
                if st.session_state.workmode == "短剧对标":
                    data = extracttemplatejson(fullresponse)
                    if data:
                        if "step4questions" in data or "step4_questions" in data or data.get("step") in ["step4ready","step4_ready"]:
                            savelabel = "Step1-3拆解分析"
                            st.session_state.benchstep4_template = data
                            st.session_state.benchstep4_confirmed = False
                        elif data.get("step") == "step7_format":
                            savelabel = "Step5-6方案大纲"
                            st.session_state.benchstep7_confirmed = False
                        elif data.get("step") == "batch_complete":
                            bi = data.get("batch_index", 1)
                            st.session_state.benchbatch_index = bi
                            st.session_state.benchtotal_batches = data.get("total_batches", 0)

                saveddir, savedname = autosave(fullresponse, sdd, savelabel)
                if savedname and not savedname.startswith("Error"):
                    fullresponse += f"\n\n---\n📁 已保存: {savedname}"
                    msgs[-1]["content"] = fullresponse

                archivecurrentchat()
                st.session_state.isgenerating = False
                st.session_state.stoprequested = False

                if '[请说继续获取下一段]' in fullresponse.replace(f"\n\n---\n📁 已保存: {savedname}", ""):
                    st.session_state.autocontinue = True

                st.rerun()

            except Exception as e:
                err_msg = str(e)[:500]
                if "timed out" in err_msg.lower() or "timeout" in err_msg.lower():
                    err_display = "请求超时。"
                elif "connection" in err_msg.lower():
                    err_display = "连接中断"
                else:
                    err_display = err_msg

                placeholder.error(f"生成失败：{err_display}")
                msgs.append({"role":"assistant","content":f"错误:{err_display}","is_error":True})
                archivecurrentchat()
                st.session_state.isgenerating = False
                st.session_state.stoprequested = False
                st.rerun()