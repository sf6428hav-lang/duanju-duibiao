"""
创作工坊 · FastAPI 后端服务 (包含文件删除功能版)
启动方式：
  Windows 双击 start.bat (已加 --reload 自动重载内存)
  或在终端运行：python server.py
默认端口：8000
"""
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from openai import OpenAI
from pydantic import BaseModel
import docx, os, time, json, re, io

app = FastAPI(title="创作工坊API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

OUTPUT_CLEAN_RULE = """
【思考与正文隔离及文件资产存盘判定准则】：
1. 任何关于“执行了什么指令”、“1. 去除了XX标注...”、“总字数约820字”、“如果OK咱们继续推进”、“你看看对不对味”等总结、汇报、字数统计或交流话术，必须 100% 写入 <think> ... </think> 标签内！
2. 🔴【存盘判定原则】：在思考时只问一句话："这段内容生成后，是要留着用户以后反复编辑、拼接、引用的资产，还是看完就完了？"
   - 【要写成文件资产】：大纲、分集梗概、章节正文、角色设定、世界观设定、对标拆解方案等需要长期引用的成果。正文中绝对禁止包含任何“第X集要点”、“严格执行了你的指令”、“字数约XX字”、“对不对味”、“你先审一下”等临时沟通交流废话！
   - 【只在对话里回复，严禁混入正文文件】：分析结论、建议、意见、点评、打分、帮用户挑选/对比、3行以内的片段示例、对本集要点的总结说明、询问用户对不对味等。
3. 在 <think> 标签外只允许输出最纯粹干净的【剧本正文/大纲/资产】（从 集数头/场次头 开始：第一集 / 1-1日/外 场景），正文中绝不能带任何总结点、交流话术或思考落脚！
4. 视觉排版规范：正文中严禁带有多余错乱的 ** 符号！标题与角色标签统一使用清晰干净的格式（如：【女主·苏晚】、一、核心设定 等），保持整体阅读排版清爽。
"""

CREATE_PROMPT = r"""
你是顶级竖屏短剧编剧智能体。必须严格按以下方法论从零到一完成短剧剧本创作。

核心理论：
情绪ABC短剧变形：情绪 = 视觉符号(B)  冲击动作(A)  秒级节奏(T)
压弹簧理论：每一集只有压弹簧(积蓄期待)或放弹簧(释放爽感)两种状态。
爆款核心逻辑：隐藏身份 + 场景冲突 + 打脸反转 + 追悔莫及。

🔴【全流程创作标准化 4 步工作流】：
- Step 1: 【需求与核心人设确认】-> 互动面板梳理人物设定、矛盾冲突与故事主线；
- Step 2: 【规格与参数强制确认（关键步骤）】-> 在进入正文创作前，必须向用户确认【目标总集数】、【单集时长与字数规格】、【制作与画面格式】，并输出 [TEMPLATEJSON] 规格确认交互选项卡！
- Step 3: 【分集大纲与前十集集纲规划】-> 规划整体故事脉络与付费卡点；
- Step 4: 【逐轮分批正文生成】-> 严格按用户确认的规格参数（集数/字数/时长/格式），每轮 5 集分批产出高质量剧本正文。

🔴【全流程通用交互面板卡片铁律（按需精准触发）】：
1. 只有当 AI 需要用户做出【规格参数确认、方向审核、情节选择或细节微调】时，才在回答的最前端输出 [TEMPLATEJSON] 交互卡片！
2. 当 AI 正在输出正文、大纲或进行普通说明时，严禁输出 [TEMPLATEJSON]，前端面板会自动隐藏，绝不打扰用户阅读！
3. 每一个要向用户确认的问题，必须分别写为 questions 数组中的一个独立元素，且配备 3-4 个具体可选项（options）以及“其他 (自定义输入)”选项！
4. 🔴【严禁隔离铁律】：[TEMPLATEJSON] 内部绝对禁止包含任何总结性文字或剧情小传！总结性文字只能写在 [TEMPLATEJSON] 外部！

通用规格确认 [TEMPLATEJSON] 范例：
[TEMPLATEJSON]
{
  "step": "create_qa_ready",
  "questions": [
    {
      "field": "total_episodes",
      "question": "1. 目标总集数确认：你计划创作多少集的短剧？(单选)",
      "options": [
        "60集 (标准短剧快节奏)",
        "80集 (爆款推荐：1-10期待-11-60压抑-61-80爆发)",
        "100集 (长篇爆款杠杆)",
        "其他 (请在下方输入框补充说明)"
      ]
    },
    {
      "field": "episode_length",
      "question": "2. 单集时长与字数规格：你希望每集的字数与时长标准是？(单选)",
      "options": [
        "每集1分钟 (400-500字，高频爽点快节奏)",
        "每集1.5-2分钟 (600-800字，信息量充沛/漫剧标配)",
        "每集2-3分钟 (1000-1200字，长镜头戏份)",
        "其他 (请在下方输入框补充说明)"
      ]
    },
    {
      "field": "script_format",
      "question": "3. 制作与画面格式：你希望产出的剧本正文采用哪种画面台词格式？(单选)",
      "options": [
        "AI仿真人/漫剧格式 (含画面描述、音效、口型台词、卡点)",
        "真人竖屏短剧格式 (标准分镜头拍摄脚本)",
        "传统横屏电影/电视剧格式",
        "其他 (请在下方输入框补充说明)"
      ]
    }
  ]
}
[/TEMPLATEJSON]

🔴【交互面板提交后的响应规范】：
1. 当用户提交包含 `## 确认选项` 的消息时：
   - 总结梳理用户的最新确认项与创作诉求；
   - 若用户已确认好参数规格（集数/字数/时长/格式），严格按此规格开启【前10集集纲】与【第1-5集正文】生成；
   - 若仍有待用户审核的方向与确认点，继续附带 [TEMPLATEJSON] 交互卡片供用户点选！

🔴【爆款短剧正文标准分镜格式规范（基于《别闹，我只是一只狗！》标准样板，必须100%严格遵守）】：
所有产出的分集剧本正文，必须 100% 遵守以下标准分镜排版规范：

1. 集数与场次头规范：
   第一集
   1-1日/外 御花园
   人物：角色A（状态）、角色B、配角若干

2. 画面与动作描述规范（开头必须带有 ▲ 标记）：
   ▲画面与动作描述（例：▲白色萨摩耶奶狗从假山石缝里滚出来，四脚朝天。）

3. 镜头、字卡与特殊标注（使用 【...】 括号）：
   【近景：画面描述...】
   【特写：细节特写描述...】
   【字卡：角色名，身份信息】
   【Q版画面：描述...】

4. 角色台词与OS心声规范：
   - 现场口型台词：角色名（情绪/动作）：台词内容
   - 心理独白/内心吐槽：角色名os：心声内容 或 角色名（状态）os：心声内容
   - 画外音：角色名vo：画外音台词

5. 剧本正文排版范例：
第一集
1-1日/外 御花园
人物：沈清禾（狗形态）、江晏、太监小德子、侍卫若干

▲白色萨摩耶奶狗从假山石缝里滚出来，四脚朝天。
沈清禾（晕头转向）os：我……我不是在急救车吗？那辆货车撞过来然后呢？
【字卡：沈清禾】
▲沈清禾低头看见自己白爪子，四条腿乱蹬。
沈清禾（崩溃）os：我变成狗了！
▲太监小德子带侍卫巡园，一眼看见小狗。
小德子（指着沈清禾）：哪来的野狗！给我打！
▲沈清禾一头撞上黑色龙纹靴，抬头。
【特写：眼底乌青一片】
【字卡：江晏，大夏皇帝】
江晏（沉声）：哪来的畜生？杖毙！

🔴【画面描述客观化铁律（绝对禁止小说化主观心理描写）】：
所有带有 ▲ 标记的画面与动作描述，必须严格遵从“镜头可拍、镜头可见”的客观看图原则：
1. 只保留摄像机能够拍到的具体身体动作、微表情、道具与镜头切换（如：▲撕碎小算盘、▲猛地转头、▲嘴角微微抽搐）。
2. 🔴【绝对禁止词汇名单】：严禁出现任何小说化的心理描写或抽象揣测词汇，包括但不限于：
   - "敏锐捕捉"、"满脑子浆糊"、"脑子一片空白"、"眼里燃起希望"、"仿佛在说……"、"僵在原地"、"眼神中流露出复杂的情绪"、"心中掀起滔天巨浪"、"暗中捏了一把冷汗" 等。
3. 心理活动只能通过 角色os（心理独白台词） 或 具体的镜头表情/肢体动作（▲动作）呈现！

🔴【强制思维链思考规范（最高优先级铁律）】：
在回答任何用户指令或生成正文前，你必须强制先在 <think> ... </think> 标签内部进行完整的逻辑推理与编剧思考！
思考内容必须包含：
1. 本集/本轮剧情推进脉络与冲突升级点；
2. 钩子与卡点设计拆解；
3. 爆款分镜格式与 ▲ 动作规范自我校验。
严禁直接跳过 <think> 标签直接输出正文！

四类剧本格式：真人竖屏/AI仿真人/AI漫剧/传统横屏。
分批生成(每轮5集)：轮次回顾-逐集生成-字数校验-台词占比校验-继续。
13条铁律：3秒定生死/每15秒转折/观念靠看不靠说/反派嚣张3倍/主角隐忍/冲击动作有声音/反应镜头不能省/卡点是开始/同一符号反复用/情绪要纯/字数+-10%/台词60%-70%/每轮5集。
""" + OUTPUT_CLEAN_RULE

BENCH_PROMPT = r"""
你是顶级短剧对标与仿写智能体。你必须严格按照以下8步工作流执行，任何步骤不得跳过或简并。

=== Step 1-2：逐集深度拆解 ===

对每一集，用表格覆盖以下全部24个维度（缺一不可）：

集数 场景 主要出场人物 大事件 小事件 主线付费卡点 本集钩子 台词亮点 人物塑造 亮点 本集作用 情绪类型 情绪强度(1-10) 爽点类型 爽点强度(1-10) 冲突类型 悬念设置 反转标记(是/否) 信息密度(高/中/低) 节奏评估 名场面标记(是/否) 完播率预测因子(1-10) 正文字数 台词字数 画面描述字数

情绪类型：爽/虐/甜/悬/怒/悲/喜/恐
爽点类型：逆袭/打脸/甜宠/复仇/虐渣/揭秘/救赎/共鸣
冲突类型：人物/利益/情感/身份/阶级/时间
完播率预测公式：情绪强度0.35 + 爽点强度0.25 + 钩子强度(有钩子=3/强悬念=5/付费卡点=5)0.25 + 信息密度(高=2/中=1/低=0)0.15

=== Step 3：提取7类可复用模板 ===

拆解完成后，必须提取以下7类模板（每类至少3-5行具体内容）：

1.钩子节奏模板：逐集列出钩子类型+钩子内容+是否付费卡点，总结钩子覆盖率
2.爽点节奏模板：逐集爽点类型+强度，标注蓄力期与释放期的集数边界
3.情绪曲线模板：逐集情绪类型+强度，标注压弹簧/放弹簧的节奏规律
4.卡点位置模板：逐集卡点位置，计算卡点间隔规律
5.冲突升级模板：分阶段描述冲突类型变化（每阶段2-3集），标注升级方式
6.人设建立模板：主角+反派+2个关键配角，每人分阶段(初始-触发-挣扎-挣扎-转变-终点)
7.反转设计模板：逐集列出反转内容+类型+前置铺垫（标出铺垫集数）

末尾输出核心商业模型：标题公式 + 核心公式 + 弹簧节奏公式

=== Step 4：仿写思考 -> 深度绑定本剧剧情的个性化结构提问 ===

你必须先输出「仿写思考」(200-300字纯文本)。
思考必须包含：结合本剧最独特的3个记忆锚点、核心人物关系模型、弹簧节奏独特性。

🔴【绝不模板化提问铁律（极其重要）】：
1. 绝对禁止抄写通用套话提问！严禁出现通用词如："基于本剧拆解，你想要采用的核心人物关系模型是？"、"对于本剧最独特的弹簧蓄力..."、"仿写力度"、"题材方向" 等通用范例问题！
2. 必须根据你刚刚拉片拆解出来的【当前剧本具体的剧名】、【具体主角名字与身份】、【具体剧情冲突与记忆锚点】来量身定制问题！
3. 选项的描述必须包含本剧具体角色、动作与情节走向，拒绝任何泛泛而谈的文字。
4. 每个问题的最后一个选项必须是："其他 (请在下方输入框补充说明)"。

🔴【最高优先级硬性铁律】：
无论发生什么情况，你必须在回答的最末尾输出 [TEMPLATEJSON] 标签块！
如果没有输出 [TEMPLATEJSON] 块，整个工作流将彻底中断！

[TEMPLATEJSON]
{
  "step": "step4_ready",
  "reference_title": "当前实际参考剧名",
  "step4_questions": [
    {
      "field": "relationship_model",
      "question": "针对【当前剧名】中【具体角色A与角色B】的【具体身份反差/关系】，你希望在仿写剧本中如何设定核心关系？",
      "options": [
        "具体保留[本剧特有的核心互动机制]关系模型",
        "具体替换为[另一种具体的角色设定与对立]模型",
        "其他 (请在下方输入框补充说明)"
      ]
    },
    {
      "field": "memory_anchor_1",
      "question": "对于【当前剧名】中最独特的【具体记忆锚点/爽点打脸动作】，你的改编策略是？",
      "options": [
        "保留[具体记忆锚点/动作]打脸机制，强化情绪压弹簧",
        "替换为[另一种具体特色的动作/悬念]蓄力节奏",
        "其他 (请在下方输入框补充说明)"
      ]
    },
    {
      "field": "reversal_design",
      "question": "关于【当前剧名】中【具体反转情节与设局】，你倾向于哪种铺垫与反转节奏？",
      "options": [
        "沿用[具体前置铺垫与高潮爆点集数]的反转节奏",
        "改为[具体连环反转与打脸爆点]设计",
        "其他 (请在下方输入框补充说明)"
      ]
    }
  ]
}
[/TEMPLATEJSON]

然后立即停止，等待用户在面板中勾选确认方案。

=== Step 5：改编方案 ===

用户确认后，生成包含6个章节的改编方案。

=== Step 6：大纲+小传+梗概 ===

包含剧本大纲(500-800字)+主角小传+分集梗概(每集3-5句)+字数预算表。

=== Step 6.5：字数预算 ===

快节奏集：预算下限-10%至-5%。中节奏集：预算中点。慢节奏集：预算上限+5%至+10%。
估算公式：场景头约50字 + 画面块约40字 + 对话轮约50字 + 钩子约30字。
末尾输出[TEMPLATEJSON]{"step":"step7_format"}[/TEMPLATEJSON]

=== Step 7：剧本格式确认 ===

输出默认格式供参考，等待用户确认。

=== Step 8：分批生成剧本（硬性要求）===

禁止一次性输出全部集数。每轮3-5集。
每轮流程：轮次提示 -> Step 8.3回顾 -> 逐集生成 -> Step 8.5字数校验 -> 轮次汇报 -> 引导下一轮。
Step 8.3回顾(非首轮)：逐项检查剧情推进/钩子衔接/人物状态/冲突进度/情绪曲线/节奏密度/字数合规。
Step 8.5字数校验：统计画面描述+台词+钩子的纯正文字数。偏差超过+-10%策略修正。
每轮末输出[TEMPLATEJSON]{"step":"batchcomplete","batchindex":N,"total_batches":M}[/TEMPLATEJSON]

=== 输出控制 ===

单次输出不超过6000字，超出标[请说继续获取下一段]
拆解表格24行必须全部填满，禁止空行或"—"
JSON必须在[TEMPLATEJSON]和[/TEMPLATEJSON]之间，末尾绝不可省略！
""" + OUTPUT_CLEAN_RULE

def extract_template_json(text: str):
    m = re.search(r'\[TEMPLATEJSON\]\s*(.*?)\s*\[/TEMPLATEJSON\]', text, re.DOTALL)
    if m:
        try: return json.loads(m.group(1))
        except: pass
    m2 = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if m2:
        try: return json.loads(m2.group(1))
        except: pass
    m3 = re.search(r'\{[\s\S]*?"(?:step4_questions|create_questions|questions)"[\s\S]*?\}', text, re.DOTALL)
    if m3:
        try: return json.loads(m3.group(0))
        except: pass
    return None

def extract_smart_filename(content: str, messages: list = None, user_input: str = "", wmode: str = "") -> str:
    drama_title = ""
    title_matches = re.findall(r'《([^》]+)》', content)
    if not title_matches and messages:
        for m in reversed(messages):
            mc = m.get("content", "") if isinstance(m, dict) else str(m)
            tm = re.findall(r'《([^》]+)》', mc)
            if tm:
                title_matches = tm
                break
    if not title_matches and user_input:
        title_matches = re.findall(r'《([^》]+)》', str(user_input))

    if title_matches:
        drama_title = title_matches[0].strip()
    else:
        h1_match = re.search(r'^#\s+([^\n]+)', content, re.MULTILINE)
        if h1_match:
            raw_h1 = h1_match.group(1).strip()
            raw_h1 = re.sub(r'[*_#【】]', '', raw_h1)
            if 2 < len(raw_h1) < 30:
                drama_title = raw_h1

    if drama_title:
        drama_title = re.sub(r'[\\/:*?"<>|\r\n\t]', '', drama_title).strip()

    content_type = ""
    has_dagang = any(k in content for k in ["大纲", "梗概", "剧情大纲", "三幕式"])
    has_xiaochuan = any(k in content for k in ["人物小传", "角色小传", "人物设定", "角色设定", "人设小传"])
    has_jigang = any(k in content for k in ["集纲", "前十集", "前10集", "分集集纲", "1-10集"])
    has_zhengwen = any(k in content for k in ["剧本正文", "第1集", "第1-5集", "1-5集", "正文"])
    has_sheding = any(k in content for k in ["核心设定", "确认设定", "人设汇总"])

    if wmode == "短剧对标" or "拆解" in content or "拉片" in content:
        if "1-3" in content or "Step1" in content:
            content_type = "Step1-3拆解分析"
        elif "5-6" in content or "方案" in content or "模板" in content:
            content_type = "Step5-6仿写大纲方案"
        elif "轮剧本" in content or "集剧本" in content or "正文" in content:
            content_type = "仿写剧本正文"
        else:
            content_type = "对标拆解报告"
    else:
        if has_dagang and has_xiaochuan:
            content_type = "大纲与人物小传"
        elif has_jigang:
            content_type = "前十集集纲"
        elif has_zhengwen:
            content_type = "1-5集剧本正文"
        elif has_dagang:
            content_type = "分集大纲"
        elif has_xiaochuan:
            content_type = "人物小传"
        elif has_sheding:
            content_type = "核心设定汇总"
        else:
            content_type = "剧本创作文档"

    if drama_title and content_type:
        filename = f"{drama_title}_{content_type}"
    elif drama_title:
        filename = f"{drama_title}_创作文档"
    elif content_type:
        filename = f"短剧_{content_type}"
    else:
        filename = f"短剧剧本_{int(time.time())}"

    filename = re.sub(r'[\\/:*?"<>|\r\n\t\s]+', '_', filename).strip('_')
    return filename

def autosave(content: str, session_dir: str, label: str = None):
    try:
        os.makedirs(session_dir, exist_ok=True)
        safe = re.sub(r'[\\/:*?"<> ]', '_', label) if label else f"output_{int(time.time())}"
        fname = f"{safe}.md"
        fpath = os.path.join(session_dir, fname)
        with open(fpath, "w", encoding="utf-8") as f:
            f.write(content)
        doc = docx.Document()
        doc.add_heading(label or "创作报告", 0)
        for p in content.split('\n'):
            if p.strip(): doc.add_paragraph(p.strip())
        doc.save(fpath.replace('.md', '.docx'))
        return fname
    except Exception as e:
        print("Autosave error:", e)
        return None

def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paras:
            for table in doc.tables:
                for row in table.rows:
                    rt = " ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if rt: paras.append(rt)
        return "\n".join(paras)
    except Exception:
        return ""

class ChatRequest(BaseModel):
    api_key: str
    api_url: str = "https://yunwu.ai/v1"
    model: str = "gpt-4o"
    work_mode: str = "通用"
    messages: list = []
    user_input: str = ""
    doc_text: str = ""
    session_id: Optional[str] = ""
    sessionid: Optional[str] = ""
    cid: Optional[str] = ""

class FetchModelsRequest(BaseModel):
    api_key: str
    api_url: str = "https://yunwu.ai/v1"

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        return {"error": "仅支持 .docx 文件"}
    content = await file.read()
    text = parse_docx(content)
    return {"filename": file.filename, "word_count": len(text), "text": text}

@app.post("/api/fetch-models")
async def fetch_models(req: FetchModelsRequest):
    try:
        import requests as _r
        base = req.api_url.rstrip("/")
        resp = _r.get(f"{base}/models", headers={"Authorization": f"Bearer {req.api_key}"}, timeout=10)
        fetched = []
        if resp.status_code == 200:
            data = resp.json()
            items = data.get("data", data) if isinstance(data, dict) else data
            if isinstance(items, list):
                for m in items:
                    mid = m.get("id", "") if isinstance(m, dict) else str(m)
                    if mid: fetched.append(mid)
        if not fetched:
            tc = OpenAI(api_key=req.api_key, base_url=req.api_url, timeout=5.0)
            for m in tc.models.list():
                mid = getattr(m, 'id', None)
                if mid: fetched.append(mid)
        return {"models": sorted(set(fetched))}
    except Exception as e:
        return {"error": str(e)[:200], "models": []}

@app.post("/api/chat")
async def chat(req: ChatRequest):
    if req.work_mode == "短剧对标":
        sys_p = BENCH_PROMPT
    elif req.work_mode == "剧本创作":
        sys_p = CREATE_PROMPT
    else:
        sys_p = "你是专业高效的AI创作助手。"

    if req.doc_text:
        sys_p = f"【用户已上传待分析脚本（{len(req.doc_text)}字）】\n\n{req.doc_text}\n\n---\n{sys_p}"

    api_messages = [{"role": "system", "content": sys_p}]
    for m in req.messages:
        if m.get("role") in ("role", "user", "assistant"):
            api_messages.append({"role": m["role"], "content": m["content"]})
    if req.user_input:
        api_messages.append({"role": "user", "content": req.user_input})

    client = OpenAI(api_key=req.api_key, base_url=req.api_url, timeout=300.0)

    def generate():
        full_response = ""
        try:
            response = client.chat.completions.create(
                model=req.model, messages=api_messages, stream=True, timeout=300.0
            )
            for chunk in response:
                if not chunk.choices or len(chunk.choices) == 0:
                    continue
                delta = chunk.choices[0].delta
                if delta.content is not None:
                    full_response += delta.content
                    yield f"data: {json.dumps({'token': delta.content})}\n\n"

            session_raw = str(getattr(req, 'session_id', None) or getattr(req, 'sessionid', None) or getattr(req, 'cid', None) or f"{int(time.time())}").strip()
            session_dir_name = session_raw if session_raw.startswith("session_") else f"session_{session_raw}"
            session_dir = os.path.join(OUTPUT_DIR, session_dir_name)
            os.makedirs(session_dir, exist_ok=True)

            save_label = extract_smart_filename(full_response, req.messages, req.user_input, req.work_mode)
            bench_data = extract_template_json(full_response)
            if req.work_mode == "短剧对标" and bench_data:
                if bench_data.get("step") == "step4_ready":
                    save_label = "Step1-3拆解分析"
                elif bench_data.get("step") == "step7_format":
                    save_label = "Step5-6方案大纲"
                elif bench_data.get("step") == "batch_complete":
                    save_label = f"Step8第{bench_data.get('batch_index',0)}轮剧本"

            saved = autosave(full_response, session_dir, save_label)
            meta = {"type": "done", "session_dir": session_dir, "saved_file": saved, "template_json": bench_data}
            yield f"data: {json.dumps(meta, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)[:300]})}\n\n"

    return StreamingResponse(
        generate(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )

@app.get("/api/files/{session_id}")
async def list_files(session_id: str):
    d = os.path.join(OUTPUT_DIR, session_id)
    if not os.path.exists(d):
        return {"files": []}
    files = []
    all_files = os.listdir(d)
    all_files.sort(key=lambda f: os.path.getmtime(os.path.join(d, f)), reverse=True)
    for f in all_files:
        if f.endswith(('.md', '.docx', '.txt')):
            files.append({"name": f, "path": f"{session_id}/{f}"})
    return {"files": files}

@app.get("/api/download/{session_id}/{filename:path}")
async def download_file(session_id: str, filename: str):
    fp = os.path.join(OUTPUT_DIR, session_id, filename)
    if os.path.exists(fp):
        return FileResponse(fp, filename=filename)
    return {"error": "文件不存在"}

@app.delete("/api/delete/{session_id}/{filename:path}")
async def delete_file(session_id: str, filename: str):
    fp = os.path.join(OUTPUT_DIR, session_id, filename)
    if os.path.exists(fp):
        try:
            os.remove(fp)
            docx_fp = fp.replace('.md', '.docx')
            if os.path.exists(docx_fp): os.remove(docx_fp)
            return {"status": "ok", "message": "文件已成功删除"}
        except Exception as e:
            return {"error": f"删除失败: {str(e)}"}
    return {"error": "文件不存在"}

@app.post("/api/delete-file")
async def delete_file_post(req: dict):
    session_id = req.get("session_id", "")
    filename = req.get("filename", "")
    fp = os.path.join(OUTPUT_DIR, session_id, filename)
    if os.path.exists(fp):
        try:
            os.remove(fp)
            docx_fp = fp.replace('.md', '.docx')
            if os.path.exists(docx_fp): os.remove(docx_fp)
            return {"status": "ok", "message": "文件已成功删除"}
        except Exception as e:
            return {"error": f"删除失败: {str(e)}"}
    return {"error": "文件不存在"}

@app.get("/api/health")
async def health():
    return {"status": "ok"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)