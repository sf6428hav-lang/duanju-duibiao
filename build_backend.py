"""运行这个脚本，自动生成 backend.py"""
import os

code = '''from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from openai import OpenAI
from pydantic import BaseModel
import docx, os, time, json, re, io

app = FastAPI(title="创作工坊API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

OUTPUT_CLEAN_RULE = """
[思考与正文隔离铁律]：
请先在 <thinking> ... </thinking> 标签内进行编剧思考与拉片拆解。
在 <thinking> 标签外输出最终正文。正文禁止带有 ** 错乱符号。
"""

CREATE_PROMPT = r"""
你是顶级竖屏短剧编剧智能体。必须严格按以下方法论从零到一完成短剧剧本创作。

核心理论：
情绪ABC短剧变形：情绪 = 视觉符号(B) 冲击动作(A) 秒级节奏(T)
压弹簧理论：每一集只有压弹簧(积蓄期待)或放弹簧(释放爽感)两种状态。
爆款核心逻辑：隐藏身份 + 场景冲突 + 打脸反转 + 追悔莫及。

创作前确认参数：总集数、单集字数(+-10%)、输出格式、台词占比(默认65%)。

模式1-头脑风暴：核心人设5要素-核心场景-情绪主轴-终极反转-付费卡点。
模式2-大纲规划：三幕式(1-10建立期待-11-60积累压抑-61-80爆发释放)。五秒节拍法。
模式3-人物小传：六维人设容器(标签/观念/视觉符号/防御机制/核心目标/最大矛盾)。反派动机必须真实。
模式4-剧本优化：冰山四维度。每场5-8句台词，有来有回层层加码。

四类剧本格式：真人竖屏/AI仿真人/AI漫剧/传统横屏。
分批生成(每轮5集)：轮次回顾-逐集生成-字数校验-台词占比校验-继续。
13条铁律：3秒定生死/每15秒转折/观念靠看不靠说/反派嚣张3倍/主角隐忍/冲击动作有声音/反应镜头不能省/卡点是开始/同一符号反复用/情绪要纯/字数+-10%/台词60%-70%/每轮5集。
""" + OUTPUT_CLEAN_RULE

BENCH_PROMPT = r"""
你是顶级短剧对标与仿写智能体。你必须严格按照以下8步工作流执行，任何步骤不得跳过或简并。

=== Step 1-2：逐集深度拆解 ===

对每一集，用表格覆盖以下全部24个维度（缺一不可）：

集数 场景 主要出场人物 大事件 小事件 主线付费卡点 本集钩子 台词亮点 人物塑造 亮点 本集作用 情绪类型 情绪强度(1-10) 爽点类型 爽点强度(1-10) 冲突类型 悬念设置 反转标记(是/否) 信息密度(高/中/低) 节奏评估 名场面标记(是/否) 完播率预测因子(1-10) 正文字数 台词字数 画面描述字数

情绪类型：爽/虐/甜/悬/怒/悲/喜/恐
爽点类型：逆袭/打脸/甜宠/复仇/虐渣/揭秘/救赎/共鸣
冲突类型：人物/利益/情感/身份/阶级/时间
完播率预测公式：情绪强度0.35 + 爽点强度0.25 + 钩子强度(有钩子=3/强悬念=5/付费卡点=5)0.25 + 信息密度(高=2/中=1/低=0)0.15

=== Step 3：提取7类可复用模板 ===

拆解完成后，必须提取以下7类模板（每类至少3-5行具体内容）：

1.钩子节奏模板：逐集列出钩子类型+钩子内容+是否付费卡点，总结钩子覆盖率
2.爽点节奏模板：逐集爽点类型+强度，标注蓄力期与释放期的集数边界
3.情绪曲线模板：逐集情绪类型+强度，标注压弹簧/放弹簧的节奏规律
4.卡点位置模板：逐集卡点位置，计算卡点间隔规律
5.冲突升级模板：分阶段描述冲突类型变化（每阶段2-3集），标注升级方式
6.人设建立模板：主角+反派+2个关键配角，每人分阶段(初始-触发-挣扎-转变-终点)
7.反转设计模板：逐集列出反转内容+类型+前置铺垫（标出铺垫集数）

末尾输出核心商业模型：标题公式 + 核心公式 + 弹簧节奏公式

在拆解正文末尾，嵌入[TEMPLATEJSON]块，包含prototypepatterns(每类模板用name和summary)、referencetitle、genre、episodecount、corehook。不用带step4_questions。

=== Step 4：仿写思考 -> 提问（严禁输出JSON问题列表）===

你必须先输出「仿写思考」再提问。思考+问题都用纯文本，禁止用JSON包裹问题。

「仿写思考」(200-300字)必须包含：
本剧最独特的3个记忆锚点——具体的名场面/视觉符号/台词，不能抽象
核心人物关系模型——谁和谁是情感引擎？
弹簧节奏的独特性——压/放的比例和强度
如果只保留一样东西到仿写中，你选什么？为什么？

然后基于思考，在正文中提出6个开放式问题。每个问题必须引用本剧的具体元素。示例：
"《白骨精》里骨妖替祁娘而活+抚养砖儿柳儿+反杀武二的核心公式，你打算怎么处理？贴身保留还是大改？"
禁止："是否保留核心设定？"

问题涵盖：仿写方向、新设定（谁替谁活）、篇幅节奏、情绪定位、差异化锚点、受众平台。

全文末尾只输出一行：[TEMPLATEJSON]{"step":"step4_ready"}[/TEMPLATEJSON]

然后立即停止，等待用户在面板中填写方案。

=== Step 5：改编方案 ===

用户确认后，生成包含6个章节的改编方案。

=== Step 6：大纲+小传+梗概 ===

包含剧本大纲(500-800字)+主角小传+分集梗概(每集3-5句)+字数预算表。

=== Step 6.5：字数预算 ===

快节奏集：预算下限-10%至-5%。中节奏集：预算中点。慢节奏集：预算上限+5%至+10%。
估算公式：场景头约50字 + 画面块约40字 + 对话轮约50字 + 钩子约30字。
末尾输出[TEMPLATEJSON]{"step":"step7_format"}[/TEMPLATEJSON]

=== Step 7：剧本格式确认 ===

输出默认格式供参考，等待用户确认。

=== Step 8：分批生成剧本（硬性要求）===

禁止一次性输出全部集数。每轮3-5集。
每轮流程：轮次提示 -> Step 8.3回顾 -> 逐集生成 -> Step 8.5字数校验 -> 轮次汇报 -> 引导下一轮。
Step 8.3回顾(非首轮)：逐项检查剧情推进/钩子衔接/人物状态/冲突进度/情绪曲线/节奏密度/字数合规。
Step 8.5字数校验：统计画面描述+台词+钩子的纯正文字数。偏差超过+-10%策略修正。
每轮末输出[TEMPLATEJSON]{"step":"batchcomplete","batchindex":N,"total_batches":M}[/TEMPLATEJSON]

=== 输出控制 ===

单次输出不超过6000字，超出标[请说继续获取下一段]
拆解表格24行必须全部填满，禁止空行或"—"
JSON必须在[TEMPLATEJSON]和[/TEMPLATEJSON]之间
""" + OUTPUT_CLEAN_RULE

def extract_template_json(text: str):
    m = re.search(r'\[TEMPLATEJSON\]\s*(.*?)\s*\[/TEMPLATEJSON\]', text, re.DOTALL)
    if m:
        try: return json.loads(m.group(1))
        except: return None
    return None

def autosave(content: str, session_dir: str, label: str = None):
    try:
        os.makedirs(session_dir, exist_ok=True)
        safe = re.sub(r'[\\/:*?"<> ]', '', label) if label else f"output_{int(time.time())}"
        fname = f"{safe}.md"
        fpath = os.path.join(session_dir, fname)
        with open(fpath, "w", encoding="utf-8") as f:
            f.write(content)
        doc = docx.Document()
        doc.add_heading("创作报告", 0)
        for p in content.split('\n'):
            if p.strip(): doc.add_paragraph(p.strip())
        doc.save(fpath.replace('.md', '.docx'))
        return fname
    except Exception:
        return None

def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paras:
            for table in doc.tables:
                for row in table.rows:
                    rt = " ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if rt: paras.append(rt)
        return "\n".join(paras)
    except Exception:
        return ""

class ChatRequest(BaseModel):
    api_key: str
    api_url: str = "https://yunwu.ai/v1"
    model: str = "gpt-4o"
    work_mode: str = "通用"
    messages: list = []
    user_input: str = ""
    doc_text: str = ""

class FetchModelsRequest(BaseModel):
    api_key: str
    api_url: str = "https://yunwu.ai/v1"

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        return {"error": "仅支持 .docx 文件"}
    content = await file.read()
    text = parse_docx(content)
    return {"filename": file.filename, "word_count": len(text), "text": text}

@app.post("/api/fetch-models")
async def fetch_models(req: FetchModelsRequest):
    try:
        import requests as _r
        base = req.api_url.rstrip("/")
        resp = _r.get(f"{base}/models", headers={"Authorization": f"Bearer {req.api_key}"}, timeout=10)
        fetched = []
        if resp.status_code == 200:
            data = resp.json()
            items = data.get("data", data) if isinstance(data, dict) else data
            if isinstance(items, list):
                for m in items:
                    mid = m.get("id", "") if isinstance(m, dict) else str(m)
                    if mid: fetched.append(mid)
        if not fetched:
            tc = OpenAI(api_key=req.api_key, base_url=req.api_url, timeout=5.0)
            for m in tc.models.list():
                mid = getattr(m, 'id', None)
                if mid: fetched.append(mid)
        return {"models": sorted(set(fetched))}
    except Exception as e:
        return {"error": str(e)[:200], "models": []}

@app.post("/api/chat")
async def chat(req: ChatRequest):
    if req.work_mode == "短剧对标":
        sys_p = BENCH_PROMPT
    elif req.work_mode == "剧本创作":
        sys_p = CREATE_PROMPT
    else:
        sys_p = "你是专业高效的AI创作助手。"

    if req.doc_text:
        sys_p = f"【用户已上传待分析脚本（{len(req.doc_text)}字）】\n\n{req.doc_text}\n\n---\n{sys_p}"

    api_messages = [{"role": "system", "content": sys_p}]
    for m in req.messages:
        if m.get("role") in ("user", "assistant"):
            api_messages.append({"role": m["role"], "content": m["content"]})
    if req.user_input:
        api_messages.append({"role": "user", "content": req.user_input})

    client = OpenAI(api_key=req.api_key, base_url=req.api_url, timeout=300.0)

    def generate():
        full_response = ""
        try:
            response = client.chat.completions.create(
                model=req.model, messages=api_messages, stream=True, timeout=300.0
            )
            for chunk in response:
                if not chunk.choices or len(chunk.choices) == 0:
                    continue
                delta = chunk.choices[0].delta
                if delta.content is not None:
                    full_response += delta.content
                    yield f"data: {json.dumps({'token': delta.content})}\n\n"

            session_dir = os.path.join(OUTPUT_DIR, f"session_{int(time.time())}")
            os.makedirs(session_dir, exist_ok=True)

            save_label = f"output_{int(time.time())}"
            bench_data = extract_template_json(full_response)
            if req.work_mode == "短剧对标" and bench_data:
                if bench_data.get("step") in ("step4_ready", "step4_ready"):
                    save_label = "Step1-3拆解分析"
                elif bench_data.get("step") == "step7_format":
                    save_label = "Step5-6方案大纲"
                elif bench_data.get("step") == "batch_complete":
                    save_label = f"Step8第{bench_data.get('batch_index',0)}轮剧本"

            saved = autosave(full_response, session_dir, save_label)
            meta = {"type": "done", "session_dir": session_dir, "saved_file": saved, "template_json": bench_data}
            yield f"data: {json.dumps(meta, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)[:300]})}\n\n"

    return StreamingResponse(
        generate(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )

@app.get("/api/files/{session_id}")
async def list_files(session_id: str):
    d = os.path.join(OUTPUT_DIR, session_id)
    if not os.path.exists(d):
        return {"files": []}
    files = []
    for f in os.listdir(d):
        if f.endswith(('.md', '.docx', '.txt')):
            files.append({"name": f, "path": f"{session_id}/{f}"})
    return {"files": files}

@app.get("/api/download/{session_id}/{filename:path}")
async def download_file(session_id: str, filename: str):
    fp = os.path.join(OUTPUT_DIR, session_id, filename)
    if os.path.exists(fp):
        return FileResponse(fp, filename=filename)
    return {"error": "文件不存在"}

@app.get("/api/health")
async def health():
    return {"status": "ok"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
'''

out = r"C:\Users\Administrator\Desktop\短剧对标\backend.py"
with open(out, "w", encoding="utf-8") as f:
    f.write(code.strip())
print(f"已生成：{out}")